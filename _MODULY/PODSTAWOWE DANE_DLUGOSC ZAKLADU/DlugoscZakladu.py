"""
PROGRAMY/DlugoscZakladu.py
Wersja: ENGINEERING_STANDARD_V32 (Button Fix)
"""

import streamlit as st
from pathlib import Path
from io import BytesIO
import sys
import math
import re

# Biblioteki do PDF
from fpdf import FPDF

# Biblioteki do DOCX
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- KONFIGURACJA ŚCIEŻEK ---
SCIEZKA_PLIKU = Path(__file__).resolve()
SCIEZKA_FOLDERU_LOKALNEGO = SCIEZKA_PLIKU.parent

SCIEZKA_BAZOWA = None
for parent in SCIEZKA_PLIKU.parents:
    if parent.name.upper() == "KALKULATORY":
        SCIEZKA_BAZOWA = parent
        break

if SCIEZKA_BAZOWA is None:
    SCIEZKA_BAZOWA = SCIEZKA_PLIKU.parents[2]

SCIEZKA_CZCIONKI = SCIEZKA_BAZOWA / "CZCIONKI"

if str(SCIEZKA_BAZOWA) not in sys.path:
    sys.path.append(str(SCIEZKA_BAZOWA))

# =============================================================================
# IMPORT DANYCH Z TABLIC
# =============================================================================

STAL_DATA_FALLBACK = {
    "B500B": 500,
    "B500A": 500,
    "B500C": 500,
    "RB500W": 500,
}

try:
    from TABLICE.ParametryStali import STEEL_TABLE
    STAL_DATA = {k: v.fyk for k, v in STEEL_TABLE.items()}
except ImportError:
    STAL_DATA = STAL_DATA_FALLBACK

BETON_DATA = {
    "C12/15":  [12, 1.6],
    "C16/20":  [16, 1.9],
    "C20/25":  [20, 2.2],
    "C25/30":  [25, 2.6],
    "C30/37":  [30, 2.9],
    "C35/45":  [35, 3.2],
    "C40/50":  [40, 3.5],
    "C45/55":  [45, 3.8],
    "C50/60":  [50, 4.1],
    "C55/67":  [55, 4.2],
    "C60/75":  [60, 4.4],
    "C70/85":  [70, 4.6],
    "C80/95":  [80, 4.8],
    "C90/105": [90, 5.0],
}

FI_LIST = [6, 8, 10, 12, 14, 16, 20, 25, 28, 32, 40]

# --- SYMBOLE UNICODE DLA PDF ---
SYM = {
    "fi": "\u03A6", "alpha": "\u03B1", "sigma": "\u03C3", "eta": "\u03B7",
    "rho": "\u03C1", "dot": "\u00B7", "bullet": "\u2022", "ge": "\u2265", "le": "\u2264",
    "ra": "\u2192"
}

# =============================================================================
# LOGIKA OBLICZENIOWA
# =============================================================================

def ObliczDlugoscZakladu(
    fi_mm: float,
    klasa_betonu: str,
    stal_nazwa: str,
    procent_naprezenia: float,
    warunki_przyczepnosci: str,
    rodzaj_preta: str,
    alfa6_proc: str,
    alfa2: float = 1.0,
    alfa3: float = 1.0,
    alfa5: float = 1.0
) -> dict:
    
    fck, fctm = BETON_DATA.get(klasa_betonu, [20, 2.2])
    fyk = STAL_DATA.get(stal_nazwa, 500)
    
    gamma_c = 1.4  # Zgodnie z PN-EN (NA)
    gamma_s = 1.15
    
    fctk_005 = 0.7 * fctm
    fctd = fctk_005 / gamma_c
    
    eta1 = 1.0 if warunki_przyczepnosci == "Dobre" else 0.7
    eta2 = 1.0
    if fi_mm > 32:
        eta2 = (132 - fi_mm) / 100.0
        
    fbd = 2.25 * eta1 * eta2 * fctd
    
    fyd = fyk / gamma_s
    sigma_sd = (procent_naprezenia / 100.0) * fyd
    lb_rqd = (fi_mm / 4.0) * (sigma_sd / fbd)
    
    alfa1 = 1.0
    alfa4 = 1.0 
    
    mapa_a6 = {"100%": 1.5, "50%": 1.4, "33%": 1.15, "25%": 1.15, "< 25%": 1.0}
    alfa6 = mapa_a6.get(alfa6_proc, 1.0)
    
    if rodzaj_preta == "Ściskany":
        alfa2 = 1.0
        alfa3 = 1.0
        alfa5 = 1.0

    alfa_global = alfa1 * alfa2 * alfa3 * alfa4 * alfa5 * alfa6
    
    warning_alfa = False
    if rodzaj_preta == "Rozciągany":
        prod_a235 = alfa2 * alfa3 * alfa5
        if prod_a235 < 0.7:
            alfa_global = alfa1 * alfa4 * alfa6 * 0.7
            warning_alfa = True
            
    l0_calc = alfa_global * lb_rqd
    l0_min_val = max(0.3 * alfa6 * lb_rqd, 15.0 * fi_mm, 200.0)
    l0_final = max(l0_calc, l0_min_val)
    
    return {
        "fi_mm": fi_mm,
        "klasa_betonu": klasa_betonu,
        "fctd": fctd,
        "stal_nazwa": stal_nazwa,
        "fyk": fyk,
        "fyd": fyd,
        "sigma_sd": sigma_sd,
        "eta1": eta1,
        "eta2": eta2,
        "fbd": fbd,
        "lb_rqd": lb_rqd,
        "alfa1": alfa1,
        "alfa2": alfa2,
        "alfa3": alfa3,
        "alfa4": alfa4,
        "alfa5": alfa5,
        "alfa6": alfa6,
        "alfa_global": alfa_global,
        "warning_alfa": warning_alfa,
        "l0_calc": l0_calc,
        "l0_min": l0_min_val,
        "l0_final": l0_final,
        "rodzaj_preta": rodzaj_preta
    }

# =============================================================================
# GENERATOR PDF
# =============================================================================

def create_pdf_report(wynik: dict, inputs: dict) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    
    font_regular = SCIEZKA_CZCIONKI / "ArialUnicode.ttf"
    font_bold = SCIEZKA_CZCIONKI / "ArialUnicode-Bold.ttf"
    
    main_font = "Arial"
    use_unicode = False
    
    if font_regular.exists():
        try:
            pdf.add_font("ArialUni", "", str(font_regular), uni=True)
            if font_bold.exists():
                pdf.add_font("ArialUni", "B", str(font_bold), uni=True)
            else:
                pdf.add_font("ArialUni", "B", str(font_regular), uni=True)
            
            pdf.add_font("ArialUni", "I", str(font_regular), uni=True)
            main_font = "ArialUni"
            use_unicode = True
        except Exception:
            use_unicode = False

    def fix_txt(text: str) -> str:
        if use_unicode:
            return str(text)
        replacements = {
            "ą": "a", "ć": "c", "ę": "e", "ł": "l", "ń": "n", "ó": "o", "ś": "s", "ź": "z", "ż": "z",
            "Ą": "A", "Ć": "C", "Ę": "E", "Ł": "L", "Ń": "N", "Ó": "O", "Ś": "S", "Ź": "S", "Ż": "Z",
            "Φ": SYM["fi"], "α": SYM["alpha"], "η": SYM["eta"], "ρ": SYM["rho"], "σ": SYM["sigma"]
        }
        text = str(text)
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    LINE_H = 6
    MARGIN_LEFT = 15
    MARGIN_TOP = 15
    X_INDENT = 20
    X_EQ = 25
    
    pdf.set_margins(MARGIN_LEFT, MARGIN_TOP, 10)
    pdf.set_auto_page_break(True, margin=15)

    def new_line(mult: float = 1.0):
        pdf.ln(LINE_H * mult)

    def w_txt(text: str, bold: bool = False, italic: bool = False, size: int = 11):
        style = ""
        if bold: style += "B"
        if italic: style += "I"
        pdf.set_font(main_font, style, size)
        pdf.write(LINE_H, fix_txt(text))

    def w_sub(text: str, size: int = 7, move_down: float = 2.0):
        orig_y, orig_x = pdf.get_y(), pdf.get_x()
        pdf.set_font(main_font, "", size)
        pdf.set_xy(orig_x, orig_y + move_down)
        pdf.write(LINE_H, fix_txt(text))
        pdf.set_xy(pdf.get_x(), orig_y)
        pdf.set_font(main_font, "", 11)

    def header_sec(text: str, num: str):
        new_line(1.2)
        pdf.set_font(main_font, "B", 12)
        pdf.cell(0, 8, fix_txt(f"{num}. {text}"), ln=True, border="B")
        new_line(0.3)

    def build_line(segments, indent=0):
        if indent > 0:
            pdf.set_x(indent)
        for item in segments:
            text = str(item[0])
            type_ = item[1]
            if type_ == 'txt': w_txt(text)
            elif type_ == 'bold': w_txt(text, bold=True)
            elif type_ == 'italic': w_txt(text, italic=True)
            elif type_ == 'sub': w_sub(text)
            elif type_ == 'sym': w_txt(text)
        new_line()

    # --- NAGŁÓWEK ---
    pdf.set_font(main_font, "B", 16)
    pdf.cell(0, 8, fix_txt("DŁUGOŚĆ ZAKŁADU PRĘTÓW ZBROJENIOWYCH"), ln=True, align="C")
    pdf.set_font(main_font, "", 10)
    pdf.cell(0, 5, fix_txt("wg PN-EN 1992-1-1"), ln=True, align="C")
    new_line(1.0)

    # 1. PARAMETRY
    header_sec("Parametry materiałowe", "1")
    build_line([("- Średnica pręta: ", 'txt'), (f"{SYM['fi']} = {wynik['fi_mm']:.0f} mm", 'txt')], X_INDENT)
    build_line([(f"- Beton: {wynik['klasa_betonu']} (f", 'txt'), ("ctd", 'sub'), (f" = {wynik['fctd']:.2f} MPa)", 'txt')], X_INDENT)
    build_line([(f"- Stal: {wynik['stal_nazwa']} (f", 'txt'), ("yk", 'sub'), (f" = {wynik['fyk']:.0f} MPa, f", 'txt'), ("yd", 'sub'), (f" = {wynik['fyd']:.1f} MPa)", 'txt')], X_INDENT)
    
    # 2. PODSTAWOWA DŁUGOŚĆ
    header_sec("Podstawowa długość zakotwienia", "2")
    build_line([(f"- Warunki przyczepności: {inputs['warunki']} (", 'txt'), (SYM['eta'], 'sym'), ("1", 'sub'), (f" = {wynik['eta1']})", 'txt')], X_INDENT)
    build_line([(f"- Współczynnik średnicy: ", 'txt'), (SYM['eta'], 'sym'), ("2", 'sub'), (f" = {wynik['eta2']:.2f}", 'txt')], X_INDENT)
    new_line(0.5)

    # fbd
    build_line([("Przyczepność graniczna: ", 'txt'), ("f", 'italic'), ("bd", 'sub'), (" = 2.25 · ", 'txt'), (SYM['eta'], 'sym'), ("1", 'sub'), (" · ", 'txt'), (SYM['eta'], 'sym'), ("2", 'sub'), (" · f", 'txt'), ("ctd", 'sub'), (f" = {wynik['fbd']:.2f} MPa", 'txt')], X_EQ)
    
    # lb,rqd
    fi_str = f"{wynik['fi_mm']:.0f}"
    build_line([("l", 'italic'), ("b,rqd", 'sub'), (" = (", 'txt'), (SYM['fi'], 'txt'), (" / 4) · (", 'txt'), (SYM['sigma'], 'sym'), ("sd", 'sub'), (" / f", 'txt'), ("bd", 'sub'), (") = ", 'txt'), (f"({fi_str} / 4) · ({wynik['sigma_sd']:.1f} / {wynik['fbd']:.2f}) = ", 'txt'), (f"{wynik['lb_rqd']:.1f} mm", 'txt')], X_EQ)

    # 3. ALFY
    header_sec(f"Współczynniki wpływu {SYM['alpha']}", "3")
    alphas = [
        (1, wynik['alfa1']),
        (2, wynik['alfa2']),
        (3, wynik['alfa3']),
        (5, wynik['alfa5']),
        (6, wynik['alfa6'], f"   ({SYM['rho']}", "1", f" = {inputs['alfa6_in']})")
    ]
    for item in alphas:
        segments = [(SYM['alpha'], 'italic'), (str(item[0]), 'sub'), (f" = {item[1]:.2f}", 'txt')]
        if len(item) > 3:
            segments.extend([(item[2], 'txt'), (item[3], 'sub'), (item[4], 'txt')])
        build_line(segments, X_EQ)

    if wynik['warning_alfa']:
        new_line(0.5)
        pdf.set_text_color(0, 0, 0)
        build_line([("Warunek EC2 8.4.4(1): ", 'txt'), (SYM['alpha'], 'italic'), ("2", 'sub'), (" · ", 'txt'), (SYM['alpha'], 'italic'), ("3", 'sub'), (" · ", 'txt'), (SYM['alpha'], 'italic'), ("5", 'sub'), (" < 0.7 -> Przyjęto iloczyn = 0.7", 'txt')], X_EQ)
        pdf.set_text_color(0, 0, 0)
    
    new_line()
    build_line([(SYM['alpha'], 'italic'), ("global", 'sub'), (f" = {wynik['alfa_global']:.2f}", 'txt')], X_EQ)

    # 4. MINIMALNA
    header_sec("Minimalna długość zakładu", "4")
    build_line([("l", 'italic'), ("0,min", 'sub'), (f" = max(0.3 · {SYM['alpha']}", 'txt'), ("6", 'sub'), (" · l", 'txt'), ("b,rqd", 'sub'), (f"; 15{SYM['fi']}; 200 mm)", 'txt')], X_EQ)
    
    val1 = 0.3 * wynik['alfa6'] * wynik['lb_rqd']
    val2 = 15.0 * wynik['fi_mm']
    build_line([("      = max(", 'txt'), (f"{val1:.1f} mm; {val2:.1f} mm; 200 mm", 'txt'), (") = ", 'txt'), (f"{wynik['l0_min']:.1f} mm", 'txt')], X_EQ)

    # 5. WYNIK
    header_sec("Obliczenie długości zakładu", "5")
    build_line([("l", 'italic'), ("0", 'sub'), (" = ", 'txt'), (SYM['alpha'], 'italic'), ("global", 'sub'), (" · l", 'txt'), ("b,rqd", 'sub'), (f" = {wynik['alfa_global']:.2f} · {wynik['lb_rqd']:.1f} = {wynik['l0_calc']:.1f} mm", 'txt')], X_EQ)
    
    new_line(2.0)
    
    # RAMKA WYNIKU
    pdf.set_fill_color(235, 235, 235)
    pdf.rect(MARGIN_LEFT, pdf.get_y(), 180, 14, 'F')
    pdf.set_y(pdf.get_y() + 4)
    pdf.set_x(MARGIN_LEFT + 5)
    
    # Etykieta: DUŻE LITERY, bez pogrubienia
    pdf.set_font(main_font, "", 12)
    pdf.write(LINE_H, fix_txt("WYMAGANA DŁUGOŚĆ ZAKŁADU:  "))
    
    # Symbol: POGRUBIONY
    pdf.set_font(main_font, "B", 12)
    pdf.write(LINE_H, fix_txt("l"))
    
    curr_x, curr_y = pdf.get_x(), pdf.get_y()
    pdf.set_xy(curr_x, curr_y + 2)
    pdf.set_font(main_font, "B", 8)
    pdf.write(LINE_H, fix_txt("0,req"))
    
    # Wynik: POGRUBIONY
    pdf.set_xy(pdf.get_x(), curr_y)
    pdf.set_font(main_font, "B", 12)
    pdf.write(LINE_H, fix_txt(f" = {wynik['l0_final']:.0f} mm"))

    return pdf.output(dest="S").encode("latin-1", "replace")

# =============================================================================
# GENERATOR DOCX
# =============================================================================

def add_docx_formatted(paragraph, text, base_bold=False):
    replacements = {
        r"\Phi": "Φ", r"\sigma": "σ", r"\alpha": "α", r"\eta": "η", 
        r"\rho": "ρ", r"\cdot": "·", r"\le": "≤", r"\ge": "≥"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    
    tokens = re.split(r'(_\{[^}]+\}|_[a-zA-Z0-9,]+)', text)
    
    for token in tokens:
        if not token: continue
        
        if token.startswith("_{") and token.endswith("}"):
            sub_text = token[2:-1]
            r = paragraph.add_run(sub_text)
            r.font.subscript = True
            if base_bold: r.bold = True
        elif token.startswith("_"):
            sub_text = token[1:]
            r = paragraph.add_run(sub_text)
            r.font.subscript = True
            if base_bold: r.bold = True
        else:
            r = paragraph.add_run(token)
            if base_bold: r.bold = True

def create_docx_report(wynik: dict, inputs: dict) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DŁUGOŚĆ ZAKŁADU PRĘTÓW ZBROJENIOWYCH")
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("wg PN-EN 1992-1-1")
    
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_line(txt, bold=False):
        p = doc.add_paragraph()
        add_docx_formatted(p, txt, base_bold=bold)

    doc.add_heading("1. Parametry materiałowe", level=1)
    add_line(f"- Średnica pręta: \\Phi = {wynik['fi_mm']:.0f} mm")
    add_line(f"- Beton: {wynik['klasa_betonu']} (f_{{ctd}} = {wynik['fctd']:.2f} MPa)")
    add_line(f"- Stal: {wynik['stal_nazwa']} (f_{{yk}} = {wynik['fyk']:.0f} MPa, f_{{yd}} = {wynik['fyd']:.1f} MPa)")

    h = doc.add_heading(level=1)
    add_docx_formatted(h, "2. Podstawowa długość zakotwienia (l_{b,rqd})")
    
    add_line(f"- Warunki przyczepności: {inputs['warunki']} (\\eta_1 = {wynik['eta1']})")
    add_line(f"- Współczynnik średnicy: \\eta_2 = {wynik['eta2']:.2f}")
    add_line(f"- Przyczepność graniczna f_{{bd}} = {wynik['fbd']:.2f} MPa")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_docx_formatted(p, f"l_{{b,rqd}} = {wynik['lb_rqd']:.1f} mm", base_bold=False)

    h = doc.add_heading(level=1)
    add_docx_formatted(h, "3. Współczynniki wpływu \\alpha")
    
    add_line(f"- \\alpha_1 = {wynik['alfa1']:.2f}")
    add_line(f"- \\alpha_2 = {wynik['alfa2']:.2f}")
    add_line(f"- \\alpha_3 = {wynik['alfa3']:.2f}")
    add_line(f"- \\alpha_5 = {wynik['alfa5']:.2f}")
    add_line(f"- \\alpha_6 = {wynik['alfa6']:.2f} (\\rho_1 = {inputs['alfa6_in']})")
    
    if wynik['warning_alfa']:
        p = doc.add_paragraph()
        r = p.add_run("Warunek EC2 8.4.4(1): ")
        r.font.color.rgb = RGBColor(0, 0, 0)
        add_docx_formatted(p, "\\alpha_2 \\cdot \\alpha_3 \\cdot \\alpha_5 < 0.7 -> Przyjęto iloczyn = 0.7")
    
    add_line(f"\\alpha_{{global}} = {wynik['alfa_global']:.2f}", bold=False)

    h = doc.add_heading(level=1)
    add_docx_formatted(h, "4. Minimalna długość zakładu (l_{0,min})")
    add_line(f"l_{{0,min}} = {wynik['l0_min']:.1f} mm")

    h = doc.add_heading(level=1)
    add_docx_formatted(h, "5. Obliczenie długości zakładu (l_0)")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_docx_formatted(p, f"l_0 = {wynik['alfa_global'] * wynik['lb_rqd']:.1f} mm")
    
    # Wynik końcowy
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    
    r = p.add_run("WYMAGANA DŁUGOŚĆ ZAKŁADU: ")
    r.bold = False
    r.font.size = Pt(12)
    
    add_docx_formatted(p, f"l_{{0,req}} = {wynik['l0_final']:.0f} mm", base_bold=True)
    p.runs[-1].font.size = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================================================================
# STRONA STREAMLIT
# =============================================================================

def StronaDlugoscZakladu():
    st.markdown(
        """
        <style>
        .block-container { padding-top: 1.5rem; padding-bottom: 1.5rem; }
        h3 { margin-top: 1.0rem !important; margin-bottom: 0.4rem !important; font-size: 1.1rem; }
        
        /* USUNIĘTO AGRESYWNY STYL DLA PRZYCISKÓW */
        
        .big-result {
            font-size: 26px; font-weight: bold; color: #2E8B57; background-color: #f0f2f6;
            padding: 15px; border-radius: 8px; text-align: center; margin-top: 20px; border: 2px solid #2E8B57;
        }
        div.row-widget.stRadio > div { flex-direction: row; gap: 16px; }
        .warning-box {
            background-color: #3e1f1f; color: #ffcccc; padding: 10px; border-radius: 5px; 
            margin-top: 10px; margin-bottom: 10px; border: 1px solid #ff4444; font-size: 14px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # TYTUŁ
    st.markdown(
        """
        <div style="text-align:center; margin-top:0.4rem; margin-bottom:0rem;">
            <span style="font-size:42px; font-weight:800; letter-spacing:1px; color:#dddddd;">
                DŁUGOŚĆ ZAKŁADU PRĘTÓW ZBROJENIOWYCH
            </span>
        </div>
        <div style="text-align:center; font-size:14px; color:#aaaaaa; margin-top:-12px; margin-bottom:0.6rem;">
            wg PN-EN 1992-1-1
        </div>
        """,
        unsafe_allow_html=True,
    )

    # 1. DANE WEJŚCIOWE
    st.markdown("### DANE WEJŚCIOWE")

    c1, c2, c3 = st.columns(3)
    with c1:
        # Default fi=12 (index 3)
        fi_mm = st.selectbox("Średnica pręta Φ [mm]", FI_LIST, index=3)
    with c2:
        klasa_betonu = st.selectbox("Klasa betonu", list(BETON_DATA.keys()), index=4)
    with c3:
        # Default B500B
        stal_opts = list(STAL_DATA.keys())
        def_idx = 0
        if "B500B" in stal_opts:
            def_idx = stal_opts.index("B500B")
        elif "B500" in stal_opts:
            def_idx = stal_opts.index("B500")
        
        stal_nazwa = st.selectbox("Klasa stali", stal_opts, index=def_idx)

    c4, c5 = st.columns([1, 2])
    with c4:
        naprezenie = st.number_input("Naprężenia w stali $\\sigma_{sd}/f_{yd}$ [%]", 0, 100, 100, 5)
    with c5:
        st.write("Warunki przyczepności")
        cr, ce = st.columns([1, 2])
        with cr:
            warunki = st.radio("wp_label", ["Dobre", "Złe"], label_visibility="collapsed")
        with ce:
            with st.expander("ℹ️ Pomoc: Warunki przyczepności (Rysunek 8.2)"):
                img_path = SCIEZKA_FOLDERU_LOKALNEGO / "DlugoscZakladu_WarunkiPrzyczepnosci.png"
                if img_path.exists(): st.image(str(img_path), use_container_width=True)

    st.markdown("---")

    # 2. WSPÓŁCZYNNIKI WPŁYWU
    st.markdown("### WSPÓŁCZYNNIKI WPŁYWU $\\alpha$")

    rodzaj_preta = st.radio("Rodzaj pręta", ["Ściskany", "Rozciągany"], index=0, horizontal=True)

    a2_val, a3_val, a5_val = 1.0, 1.0, 1.0

    if rodzaj_preta == "Rozciągany":
        col_a2, col_a3, col_a5 = st.columns(3)
        
        with col_a2:
            st.markdown("""<b>$\\alpha_2$: Wpływ otuliny ($c_d$)</b>""", unsafe_allow_html=True)
            st.write("Czy uwzględnić wpływ otuliny?")
            u_a2_str = st.radio("a2_yn", ["Tak", "Nie"], index=1, label_visibility="collapsed", horizontal=True)
            u_a2 = (u_a2_str == "Tak")
            
            cd_in = 30.0
            if u_a2:
                with st.expander("ℹ️ Pomoc: Rysunek $c_d$"):
                    img_cd = SCIEZKA_FOLDERU_LOKALNEGO / "DlugoscZakladu_Wspolczynnik cd.png"
                    if img_cd.exists(): st.image(str(img_cd), use_container_width=True)
                cd_in = st.number_input("Współczynnik $c_d$ [mm]", value=30.0, step=1.0)
                
            val = 1.0 - 0.15 * (cd_in - fi_mm) / fi_mm
            if not u_a2: val = 1.0
            a2_val = max(0.7, min(1.0, val))

        with col_a3:
            st.markdown("""<b>$\\alpha_3$: Zbrojenie poprzeczne (nieprzyspojone)</b>""", unsafe_allow_html=True)
            st.write("Czy uwzględnić zbrojenie poprzeczne?")
            u_a3_str = st.radio("a3_yn", ["Tak", "Nie"], index=1, label_visibility="collapsed", horizontal=True)
            u_a3 = (u_a3_str == "Tak")
            
            K_in, sum_ast_in, sum_ast_min_in = 0.05, 0.0, 2.5
            if u_a3:
                with st.expander("ℹ️ Pomoc: Rysunek $K$"):
                    img_k = SCIEZKA_FOLDERU_LOKALNEGO / "DlugoscZakladu_Wspolczynnik K.png"
                    if img_k.exists(): st.image(str(img_k), use_container_width=True)
                K_in = st.selectbox("Współczynnik $K$", [0.1, 0.05, 0.0], index=1)
                sum_ast_in = st.number_input("$\\Sigma A_{st}$ [cm²]", value=0.0, step=0.1)
                sum_ast_min_in = st.number_input("$\\Sigma A_{st,min}$ [cm²]", value=2.5, step=0.1)
            
            As_1 = (math.pi * fi_mm**2)/400.0
            val_a3 = 1.0
            if u_a3 and As_1 > 0:
                lamb = (sum_ast_in - sum_ast_min_in) / As_1
                val_a3 = 1.0 - K_in * lamb
            a3_val = max(0.7, min(1.0, val_a3))

        with col_a5:
            st.markdown("""<b>$\\alpha_5$: Nacisk poprzeczny</b>""", unsafe_allow_html=True)
            st.write("Czy uwzględnić nacisk poprzeczny?")
            u_a5_str = st.radio("a5_yn", ["Tak", "Nie"], index=1, label_visibility="collapsed", horizontal=True)
            u_a5 = (u_a5_str == "Tak")
            
            p_in = 8.0 
            if u_a5:
                st.write("p - nacisk poprzeczny w [MPa] wzdłuż lbd w stanie granicznym nośności")
                p_in = st.number_input("p_input_label", value=8.0, step=0.5, label_visibility="collapsed")
            
            val_a5 = 1.0
            if u_a5:
                val_a5 = 1.0 - 0.04 * p_in
            a5_val = max(0.7, min(1.0, val_a5))

    st.markdown("<br>", unsafe_allow_html=True)

    # ALFA 6
    st.markdown("#### $\\alpha_6$: Udział prętów łączonych w przekroju ($\\rho_1$)")
    
    alfa6_proc = st.radio(
        "a6_label",
        ["100%", "50%", "33%", "25%", "< 25%"],
        index=1, # Default 50%
        horizontal=True,
        label_visibility="collapsed"
    )
    
    with st.expander("ℹ️ Pomoc: Wartości współczynników $\\alpha$ (Tablica 8.2 i 8.3)"):
        c_h1, c_h2 = st.columns(2)
        with c_h1:
            img_a15 = SCIEZKA_FOLDERU_LOKALNEGO / "DlugoscZakladu_alfa1-alfa5.png"
            if img_a15.exists(): st.image(str(img_a15), use_container_width=True)
        with c_h2:
            img_a6 = SCIEZKA_FOLDERU_LOKALNEGO / "DlugoscZakladu_alfa6.png"
            if img_a6.exists(): st.image(str(img_a6), use_container_width=True)

    st.markdown("---")

    # PRZYCISK - DODANO TYPE PRIMARY
    _, c_btn, _ = st.columns([1, 2, 1])
    with c_btn:
        oblicz = st.button("OBLICZ DŁUGOŚĆ ZAKŁADU", type="primary", use_container_width=True)

    if oblicz:
        wynik = ObliczDlugoscZakladu(
            fi_mm=float(fi_mm),
            klasa_betonu=klasa_betonu,
            stal_nazwa=stal_nazwa,
            procent_naprezenia=float(naprezenie),
            warunki_przyczepnosci=warunki,
            rodzaj_preta=rodzaj_preta,
            alfa6_proc=alfa6_proc,
            alfa2=a2_val,
            alfa3=a3_val,
            alfa5=a5_val
        )
        
        st.session_state["wynik_dl"] = wynik
        st.session_state["inputs_dl"] = {
            "naprezenie": naprezenie,
            "warunki": warunki,
            "alfa6_in": alfa6_proc
        }
        st.session_state["pokaz_dl"] = True

    # WYNIKI
    if st.session_state.get("pokaz_dl", False):
        res = st.session_state["wynik_dl"]
        inp = st.session_state["inputs_dl"]
        
        st.markdown(
            f"""
            <div class="big-result">
                Długość zakładu <i>l</i><sub>0</sub> = {res['l0_final']:.0f} mm
            </div>
            """,
            unsafe_allow_html=True,
        )
        
        st.markdown("<div style='height: 20px;'></div>", unsafe_allow_html=True)
        
        col_pdf, col_docx = st.columns(2)
        with col_pdf:
            data_pdf = create_pdf_report(res, inp)
            st.download_button(
                "📄 POBIERZ RAPORT PDF",
                data_pdf,
                file_name="DlugoscZakladu.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        with col_docx:
            data_docx = create_docx_report(res, inp)
            st.download_button(
                "📝 POBIERZ RAPORT WORD",
                data_docx,
                file_name="DlugoscZakladu.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with st.expander("Szczegóły obliczeń", expanded=False):
            
            st.markdown("#### 1. Parametry materiałowe")
            st.write(f"Średnica pręta: **Φ = {res['fi_mm']:.0f} mm**")
            st.write(f"Beton: **{res['klasa_betonu']}** ($f_{{ctd}} = {res['fctd']:.2f}$ MPa)")
            st.write(f"Stal: **{res['stal_nazwa']}** ($f_{{yk}} = {res['fyk']:.0f}$ MPa, $f_{{yd}} = {res['fyd']:.1f}$ MPa)")
            
            st.markdown("#### 2. Podstawowa długość zakotwienia ($l_{b,rqd}$)")
            st.write(f"Warunki przyczepności: **{inp['warunki']}** ($\\eta_1 = {res['eta1']}$)")
            st.write(f"Współczynnik średnicy: $\\eta_2 = {res['eta2']:.2f}$")
            st.write(f"Przyczepność graniczna $f_{{bd}}$:")
            
            st.latex(rf"f_{{bd}} = 2.25 \cdot \eta_1 \cdot \eta_2 \cdot f_{{ctd}} = 2.25 \cdot {res['eta1']} \cdot {res['eta2']:.2f} \cdot {res['fctd']:.2f} = \mathbf{{{res['fbd']:.2f}}} \text{{ MPa}}")
            st.latex(rf"l_{{b,rqd}} = \frac{{\Phi}}{{4}} \cdot \frac{{\sigma_{{sd}}}}{{f_{{bd}}}} = \frac{{{res['fi_mm']:.0f}}}{{4}} \cdot \frac{{{res['sigma_sd']:.1f}}}{{{res['fbd']:.2f}}} = \mathbf{{{res['lb_rqd']:.1f}}} \text{{ mm}}")
            
            st.markdown("#### 3. Współczynniki wpływu $\\alpha$")
            st.write(f"$\\alpha_1 = {res['alfa1']:.2f}$")
            st.write(f"$\\alpha_2 = {res['alfa2']:.2f}$")
            st.write(f"$\\alpha_3 = {res['alfa3']:.2f}$")
            st.write(f"$\\alpha_5 = {res['alfa5']:.2f}$")
            st.write(f"$\\alpha_6 = {res['alfa6']:.2f}$ ($\\rho_1$ = {inp['alfa6_in']})")
            
            if res['warning_alfa']:
                st.markdown(
                    f"""
                    <div class="warning-box">
                    ⚠️ Warunek EC2 8.4.4(1): $\\alpha_2 \\cdot \\alpha_3 \\cdot \\alpha_5 = {res['alfa2']*res['alfa3']*res['alfa5']:.3f} < 0.7$
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
                st.latex(r"\rightarrow \text{Przyjęto: } \alpha_2 \cdot \alpha_3 \cdot \alpha_5 = 0.7")
            
            st.latex(rf"\alpha_{{global}} = \alpha_1 \cdot \alpha_2 \cdot \alpha_3 \cdot \alpha_5 \cdot \alpha_6 = {res['alfa_global']:.2f}")
            
            st.markdown("#### 4. Minimalna długość zakładu ($l_{0,min}$)")
            val1 = 0.3 * res['alfa6'] * res['lb_rqd']
            val2 = 15.0 * res['fi_mm']
            st.latex(rf"l_{{0,min}} = \max(0.3 \cdot \alpha_6 \cdot l_{{b,rqd}}; 15\Phi; 200) = \max({val1:.1f} \text{{ mm}}; {val2:.1f} \text{{ mm}}; 200 \text{{ mm}}) = {res['l0_min']:.1f} \text{{ mm}}")
            
            st.markdown("#### 5. Obliczenie długości zakładu ($l_0$)")
            l0_calc = res['alfa_global'] * res['lb_rqd']
            st.latex(rf"l_0 = \alpha_{{global}} \cdot l_{{b,rqd}} = {res['alfa_global']:.2f} \cdot {res['lb_rqd']:.1f} = {l0_calc:.1f} \text{{ mm}}")
            st.latex(rf"l_{{0,req}} = \max(l_0; l_{{0,min}}) = \mathbf{{{res['l0_final']:.1f}}} \text{{ mm}}")

if __name__ == "__main__":
    StronaDlugoscZakladu()