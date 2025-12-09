"""
PROGRAMY/OtulinaZbrojenia.py
Wersja: ENGINEERING_STANDARD_V57 (Button Fix)
"""

import streamlit as st
from pathlib import Path
from io import BytesIO
import sys

# Biblioteki do PDF
from fpdf import FPDF

# Biblioteki do DOCX
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- KONFIGURACJA ŚCIEŻEK ---
SCIEZKA_PLIKU = Path(__file__).resolve()
SCIEZKA_FOLDERU_LOKALNEGO = SCIEZKA_PLIKU.parent

SCIEZKA_BAZOWA = None
for parent in SCIEZKA_PLIKU.parents:
    if parent.name.upper() == "KALKULATORY":
        SCIEZKA_BAZOWA = parent
        break

if SCIEZKA_BAZOWA is None:
    SCIEZKA_BAZOWA = SCIEZKA_PLIKU.parents[2]

SCIEZKA_CZCIONKI = SCIEZKA_BAZOWA / "CZCIONKI"

if str(SCIEZKA_BAZOWA) not in sys.path:
    sys.path.append(str(SCIEZKA_BAZOWA))

try:
    from TABLICE.ParametryBetonu import list_concrete_classes
except ImportError:
    pass


# =============================================================================
# MODUŁ OBLICZENIOWY
# =============================================================================

def get_structural_class_adjustment(
    klasa_betonu: str,
    klasa_ekspozycji: str,
    zywotnosc_100_lat: bool,
    element_plytowy: bool,
    kontrola_jakosci: bool,
) -> int:
    zmiana = 0

    # +2 za okres 100 lat
    if zywotnosc_100_lat:
        zmiana += 2

    # fck z klasy betonu
    try:
        fck = int(klasa_betonu.split("/")[0].replace("C", ""))
    except Exception:
        fck = 20

    redukcja_wytrzymalosc = False

    # Kryteria z Tablicy 4.3N
    if klasa_ekspozycji == "XC1" and fck >= 30:
        redukcja_wytrzymalosc = True
    elif klasa_ekspozycji in ["XC2", "XC3"] and fck >= 35:
        redukcja_wytrzymalosc = True
    elif klasa_ekspozycji == "XC4" and fck >= 40:
        redukcja_wytrzymalosc = True
    elif klasa_ekspozycji in ["XD1", "XD2", "XS1"] and fck >= 40:
        redukcja_wytrzymalosc = True
    elif klasa_ekspozycji in ["XD3", "XS2", "XS3"] and fck >= 45:
        redukcja_wytrzymalosc = True

    if redukcja_wytrzymalosc:
        zmiana -= 1

    # -1 za płytę w XC1
    if element_plytowy and klasa_ekspozycji == "XC1":
        zmiana -= 1

    # -1 za specjalną kontrolę
    if kontrola_jakosci:
        zmiana -= 1

    return zmiana


def get_c_min_dur_value(klasa_ekspozycji: str, klasa_konstrukcji: str) -> float:
    # Tablica 4.4N (pręty zbrojeniowe)
    mapa_S = {"S1": 0, "S2": 1, "S3": 2, "S4": 3, "S5": 4, "S6": 5}
    idx = mapa_S.get(klasa_konstrukcji, 3)

    tabela_4_4N = [
        [10, 10, 10, 15, 20, 25, 30],  # S1
        [10, 10, 15, 20, 25, 30, 35],  # S2
        [10, 10, 20, 25, 30, 35, 40],  # S3
        [10, 15, 25, 30, 35, 40, 45],  # S4
        [15, 20, 30, 35, 40, 45, 50],  # S5
        [20, 25, 35, 40, 45, 55, 55],  # S6
    ]

    mapa_col = {
        "X0": 0,
        "XC1": 1,
        "XC2": 2,
        "XC3": 2,
        "XC4": 3,
        "XD1": 4,
        "XS1": 4,
        "XD2": 5,
        "XS2": 5,
        "XD3": 6,
        "XS3": 6,
    }

    col_idx = mapa_col.get(klasa_ekspozycji, 1)
    return float(tabela_4_4N[idx][col_idx])


def ObliczOtuline(
    klasa_ekspozycji: str,
    fi_mm: float,
    klasa_betonu: str,
    zywotnosc_100_lat: bool,
    element_plytowy: bool,
    kontrola_jakosci: bool,
    beton_na_gruncie: str,
    dg_gt_32: bool,
    delta_dur_gamma: float,
    delta_dur_st: float,
    delta_dur_add: float,
    delta_dev: float,
) -> dict:
    zmiana = get_structural_class_adjustment(
        klasa_betonu,
        klasa_ekspozycji,
        zywotnosc_100_lat,
        element_plytowy,
        kontrola_jakosci,
    )
    final_num = max(1, min(6, 4 + zmiana))
    klasa_konstrukcji_final = f"S{final_num}"

    c_min_dur = get_c_min_dur_value(klasa_ekspozycji, klasa_konstrukcji_final)

    c_min_b = fi_mm
    if dg_gt_32:
        c_min_b += 5.0

    c_min_dur_mod = c_min_dur + delta_dur_gamma - delta_dur_st - delta_dur_add
    c_min = max(c_min_b, c_min_dur_mod, 10.0)

    c_nom = c_min + delta_dev

    k1 = 40.0
    k2 = 75.0
    limit_gruntu = 0.0
    if beton_na_gruncie == "Na przygotowanym podłożu":
        limit_gruntu = k1
    elif beton_na_gruncie == "Bezpośrednio na gruncie":
        limit_gruntu = k2

    c_nom = max(c_nom, limit_gruntu)

    return {
        "klasa_ekspozycji": klasa_ekspozycji,
        "klasa_konstrukcji_final": klasa_konstrukcji_final,
        "c_min_dur": c_min_dur,
        "c_min_b": c_min_b,
        "c_min": c_min,
        "c_nom": c_nom,
        "zmiana_klasy": zmiana,
        "limit_gruntu": limit_gruntu,
    }


# =============================================================================
# GENERATOR PDF (RAPORT INŻYNIERSKI)
# =============================================================================

def create_pdf_report(wynik: dict, inputs: dict) -> bytes:
    pdf = FPDF()
    pdf.add_page()

    font_regular = SCIEZKA_CZCIONKI / "ArialUnicode.ttf"
    font_bold = SCIEZKA_CZCIONKI / "ArialUnicode-Bold.ttf"

    use_unicode = False
    main_font = "Arial"

    if font_regular.exists():
        try:
            pdf.add_font("ArialUni", "", str(font_regular), uni=True)
            if font_bold.exists():
                pdf.add_font("ArialUni", "B", str(font_bold), uni=True)
            else:
                pdf.add_font("ArialUni", "B", str(font_regular), uni=True)
            pdf.add_font("ArialUni", "I", str(font_regular), uni=True)
            main_font = "ArialUni"
            use_unicode = True
        except Exception:
            use_unicode = False

    def fix_txt(text: str) -> str:
        if use_unicode:
            return str(text)
        replacements = {
            "ą": "a", "ć": "c", "ę": "e", "ł": "l", "ń": "n", "ó": "o", "ś": "s", "ź": "z", "ż": "z",
            "Ą": "A", "Ć": "C", "Ę": "E", "Ł": "L", "Ń": "N", "Ó": "O", "Ś": "S", "Ź": "S", "Ż": "Z",
        }
        text = str(text)
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    LINE_H = 5.5
    MARGIN_LEFT = 15
    MARGIN_TOP = 15

    pdf.set_margins(MARGIN_LEFT, MARGIN_TOP, 10)
    pdf.set_auto_page_break(True, margin=15)

    def new_line(mult: float = 1.0):
        pdf.ln(LINE_H * mult)

    def w_txt(text: str, bold: bool = False, size: int = 10):
        style = "B" if bold else ""
        pdf.set_font(main_font, style, size)
        pdf.write(LINE_H, fix_txt(text))

    def w_txt_ln(text: str, bold: bool = False, size: int = 10):
        w_txt(text, bold, size)
        new_line()

    def w_symbol_sub(main_char: str, sub_char: str, value_str: str = "", bold: bool = False):
        pdf.set_font(main_font, "B" if bold else "", 10)
        pdf.write(LINE_H, fix_txt(main_char))

        x = pdf.get_x()
        y = pdf.get_y()
        pdf.set_xy(x, y + 1.5)
        pdf.set_font(main_font, "B" if bold else "", 7)
        pdf.write(LINE_H, fix_txt(sub_char))

        new_x = pdf.get_x()
        pdf.set_xy(new_x, y)
        pdf.set_font(main_font, "B" if bold else "", 10)
        if value_str:
            pdf.write(LINE_H, fix_txt(value_str))

    def header_sec(text: str, num: str):
        new_line(1.2)
        pdf.set_font(main_font, "B", 11)
        pdf.cell(0, 7, fix_txt(f"{num}. {text}"), ln=True, border="B")
        new_line(0.3)

    def add_space_before_mm(text: str) -> str:
        text = str(text)
        text = text.replace("mm", " mm")
        text = text.replace("  mm", " mm")
        return text

    # --- NAGŁÓWEK ---
    pdf.set_font(main_font, "B", 16)
    pdf.cell(0, 8, fix_txt("OTULINA PRĘTÓW ZBROJENIOWYCH"), ln=True, align="C")
    pdf.set_font(main_font, "", 10)
    pdf.cell(0, 5, fix_txt("wg PN-EN 1992-1-1"), ln=True, align="C")
    new_line(1.0)

    # 1. DANE PROJEKTOWE
    header_sec("Dane projektowe i założenia", "1")

    pdf.set_font(main_font, "", 10)
    col_w = 90

    def row_data(lbl, val):
        pdf.cell(col_w, LINE_H, fix_txt(lbl), border=0)
        if "d_g" in val:
            clean_val = val.replace("$", "").replace("d_g", "").strip()
            clean_val = add_space_before_mm(clean_val)
            pdf.set_font(main_font, "", 10)
            pdf.write(LINE_H, fix_txt("d"))
            curr_x = pdf.get_x()
            curr_y = pdf.get_y()
            pdf.set_xy(curr_x, curr_y + 1.5)
            pdf.set_font(main_font, "", 7)
            pdf.write(LINE_H, fix_txt("g"))
            pdf.set_xy(pdf.get_x(), curr_y)
            pdf.set_font(main_font, "", 10)
            pdf.write(LINE_H, fix_txt(f" {clean_val}"))
            pdf.ln()
            pdf.set_font(main_font, "", 10)
        else:
            pdf.set_font(main_font, "", 10)
            val = add_space_before_mm(val)
            pdf.cell(0, LINE_H, fix_txt(val), border=0, ln=1)
            pdf.set_font(main_font, "", 10)

    row_data("Klasa ekspozycji:", inputs["klasa_ekspozycji"])
    row_data("Klasa betonu:", inputs["klasa_betonu"])
    row_data("Średnica pręta zbrojeniowego:", f'{inputs["fi_mm"]} mm')
    row_data("Uziarnienie kruszywa:", inputs["kruszywo_opis"])

    if inputs.get("betonowanie_grunt") != "Nie":
        row_data("Betonowanie na gruncie:", inputs["betonowanie_grunt"])

    # 2. KLASYFIKACJA
    header_sec("Klasa konstrukcji", "2")
    w_txt_ln("Bazowa klasa konstrukcji: S4")
    w_txt_ln("Modyfikacja klasy konstrukcji ze względu na warunki projektowe:")

    zmiana = wynik["zmiana_klasy"]
    if zmiana == 0:
        w_txt_ln("- Brak modyfikacji klasy konstrukcji.", size=10)
    else:
        if inputs["zywotnosc_100"]:
            w_txt_ln("    ΔS = +2: Zwiększenie okresu użytkowania (100 lat)")

        try:
            fck = int(inputs["klasa_betonu"].split("/")[0].replace("C", ""))
        except Exception:
            fck = 20
        eksp = inputs["klasa_ekspozycji"]
        red_wytrz = False
        if eksp == "XC1" and fck >= 30:
            red_wytrz = True
        elif eksp in ["XC2", "XC3"] and fck >= 35:
            red_wytrz = True
        elif eksp == "XC4" and fck >= 40:
            red_wytrz = True
        elif eksp in ["XD1", "XD2", "XS1"] and fck >= 40:
            red_wytrz = True
        elif eksp in ["XD3", "XS2", "XS3"] and fck >= 45:
            red_wytrz = True

        if red_wytrz:
            w_txt_ln(
                f"    ΔS = -1: Redukcja ze względu na klasę wytrzymałości ({inputs['klasa_betonu']})"
            )

        if inputs["plyta"] and eksp == "XC1":
            w_txt_ln("    ΔS = -1: Element płytowy w klasie ekspozycji XC1")

        if inputs["kontrola_jakosci_str"] == "Tak":
            w_txt_ln("    ΔS = -1: Specjalna kontrola jakości produkcji betonu")

    new_line(0.5)
    w_txt("Obliczeniowa klasa konstrukcji: ")
    w_txt(f"{wynik['klasa_konstrukcji_final']}", bold=True)
    new_line()

    # 3. WYZNACZENIE OTULINY
    header_sec("Wyznaczenie otuliny minimalnej", "3")

    w_txt("Ze względu na przyczepność: ", bold=False)
    w_symbol_sub("c", "min,b", bold=False)
    new_line()
    w_txt("     Wymagane: ")
    w_symbol_sub("c", "min,b", f" ≥ {inputs['fi_mm']} mm")
    if inputs["kruszywo_opis"] == "d_g > 32 mm":
        w_txt(" (+ 5 mm ze względu na ")
        w_symbol_sub("d", "g", " > 32 mm)")
    new_line()
    w_txt("     Przyjęto: ")
    w_symbol_sub("c", "min,b", f" = {wynik['c_min_b']:.0f} mm", bold=True)
    new_line(1.5)

    w_txt("Ze względu na warunki środowiskowe: ", bold=False)
    w_symbol_sub("c", "min,dur", bold=False)
    new_line()
    w_txt(
        f"     Dla klasy ekspozycji {inputs['klasa_ekspozycji']} i klasy konstrukcji {wynik['klasa_konstrukcji_final']} wg Tablicy 4.4N:"
    )
    new_line()
    w_txt("     Wartość bazowa: ")
    w_symbol_sub("c", "min,dur", f" = {wynik['c_min_dur']:.0f} mm", bold=True)
    new_line(1.5)

    w_txt("Całkowita otulina minimalna ", bold=False)
    w_symbol_sub("c", "min", bold=False)
    new_line()
    w_txt("     Wzór normowy: ")
    w_symbol_sub("c", "min", " = max( ")
    w_symbol_sub("c", "min,b", "; ")
    w_symbol_sub("c", "min,dur", " + ")
    w_symbol_sub("Δc", "dur,γ", " - ")
    w_symbol_sub("Δc", "dur,st", " - ")
    w_symbol_sub("Δc", "dur,add", "; 10 mm )")
    new_line()

    w_txt("     Obliczenie: ")
    w_symbol_sub(
        "c",
        "min",
        f" = max( {wynik['c_min_b']:.0f} mm; {wynik['c_min_dur']:.0f} mm + {inputs['dc_gamma']:.0f} mm - {inputs['dc_st']:.0f} mm - {inputs['dc_add']:.0f} mm; 10 mm )",
    )
    new_line()
    w_txt("     Wynik: ")
    w_symbol_sub("c", "min", f" = {wynik['c_min']:.0f} mm", bold=True)
    new_line()

    # 4. OTULINA NOMINALNA
    header_sec("Wyznaczenie otuliny nominalnej", "4")

    w_txt("Otulinę nominalną wyznacza się przez dodanie odchyłki wykonawczej: ")
    w_symbol_sub("c", "nom", " = ")
    w_symbol_sub("c", "min", " + ")
    w_symbol_sub("Δc", "dev")
    new_line(1.5)

    w_txt(f"Przyjęta odchyłka wykonawcza: ")
    w_symbol_sub("Δc", "dev", f" = {inputs['delta_dev']} mm")
    new_line()

    if wynik["limit_gruntu"] > 0:
        w_txt(f"Uwzględniono warunek betonowania na gruncie: ")
        w_symbol_sub("c", "nom", f" ≥ {wynik['limit_gruntu']:.0f} mm")
        new_line()

    new_line(0.5)
    w_txt("Obliczona wartość: ")
    val = wynik["c_min"] + float(inputs["delta_dev"])
    if wynik["limit_gruntu"] > val:
        val = wynik["limit_gruntu"]

    w_symbol_sub("c", "nom", f" = {val:.0f} mm", bold=True)
    new_line(2.0)

    # 5. PODSUMOWANIE
    pdf.set_fill_color(235, 235, 235)
    pdf.rect(MARGIN_LEFT, pdf.get_y(), 180, 12, "F")
    pdf.set_y(pdf.get_y() + 3)
    pdf.set_x(MARGIN_LEFT + 5)

    # WIELKIE LITERY, BEZ POGRUBIENIA
    pdf.set_font(main_font, "", 12)
    w_txt("WYMAGANA OTULINA NOMINALNA:  ")
    
    # SYMBOL I WYNIK POGRUBIONE
    w_symbol_sub("c", "nom", f" = {wynik['c_nom']:.0f} mm", bold=True)

    return pdf.output(dest="S").encode("latin-1", "replace")


# =============================================================================
# GENERATOR DOCX (RAPORT INŻYNIERSKI)
# =============================================================================


def create_docx_report(wynik: dict, inputs: dict) -> BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Times New Roman"
        h1.font.size = Pt(12)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_after = Pt(6)
        h1.paragraph_format.space_before = Pt(12)

    def add_p(text="", bold=False, indent=0, space_after=6):
        p = doc.add_paragraph()
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.font.bold = bold
        return p

    def run(p, text, bold=False, sub=False):
        r = p.add_run(text)
        r.font.bold = bold
        if sub:
            r.font.subscript = True
        return r

    def add_space_before_mm(text: str) -> str:
        text = str(text)
        text = text.replace("mm", " mm")
        text = text.replace("  mm", " mm")
        return text

    # Tytuł
    p = add_p()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "OTULINA PRĘTÓW ZBROJENIOWYCH", bold=True).font.size = Pt(16)
    p = add_p()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "wg PN-EN 1992-1-1")

    doc.add_paragraph("_" * 75).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Dane projektowe
    doc.add_heading("1. Dane projektowe i założenia", level=1)

    def list_item(label, val):
        p = add_p(indent=0.5, space_after=2)
        run(p, label + " ")
        val_str = add_space_before_mm(val)
        if "d_g" in val_str:
            clean_val = val_str.replace("$", "").replace("d_g", "").strip()
            clean_val = add_space_before_mm(clean_val)
            run(p, "d", bold=False)
            run(p, "g", bold=False, sub=True)
            run(p, f" {clean_val}", bold=False)
        else:
            run(p, val_str, bold=False)

    list_item("Klasa ekspozycji:", inputs["klasa_ekspozycji"])
    list_item("Klasa betonu:", inputs["klasa_betonu"])
    list_item("Średnica pręta:", f"{inputs['fi_mm']} mm")
    list_item("Uziarnienie kruszywa:", inputs["kruszywo_opis"])

    if inputs.get("betonowanie_grunt") != "Nie":
        list_item("Betonowanie na gruncie:", inputs["betonowanie_grunt"])

    # 2. Klasyfikacja
    doc.add_heading("2. Klasa konstrukcji", level=1)
    p = add_p("Bazowa klasa konstrukcji: S4", space_after=6)
    p = add_p("Modyfikacja klasy konstrukcji ze względu na warunki projektowe:", space_after=2)

    zmiana = wynik["zmiana_klasy"]
    if zmiana == 0:
        add_p("- Brak modyfikacji.", indent=1.0)
    else:
        if inputs["zywotnosc_100"]:
            add_p("ΔS = +2: Zwiększenie okresu użytkowania (100 lat)", indent=1.0)

        try:
            fck = int(inputs["klasa_betonu"].split("/")[0].replace("C", ""))
        except Exception:
            fck = 20
        eksp = inputs["klasa_ekspozycji"]
        red_wytrz = False
        if eksp == "XC1" and fck >= 30:
            red_wytrz = True
        elif eksp in ["XC2", "XC3"] and fck >= 35:
            red_wytrz = True
        elif eksp == "XC4" and fck >= 40:
            red_wytrz = True
        elif eksp in ["XD1", "XD2", "XS1"] and fck >= 40:
            red_wytrz = True
        elif eksp in ["XD3", "XS2", "XS3"] and fck >= 45:
            red_wytrz = True

        if red_wytrz:
            add_p(
                f"ΔS = -1: Redukcja ze względu na klasę wytrzymałości ({inputs['klasa_betonu']})",
                indent=1.0,
            )

        if inputs["plyta"] and eksp == "XC1":
            add_p("ΔS = -1: Element płytowy w klasie ekspozycji XC1", indent=1.0)

        if inputs["kontrola_jakosci_str"] == "Tak":
            add_p("ΔS = -1: Specjalna kontrola jakości produkcji betonu", indent=1.0)

    p = add_p(space_after=12)
    p.paragraph_format.space_before = Pt(6)
    run(p, "Obliczeniowa klasa konstrukcji: ")
    run(p, f"{wynik['klasa_konstrukcji_final']}", bold=True)

    # 3. Otulina minimalna
    doc.add_heading("3. Wyznaczenie otuliny minimalnej", level=1)

    p = add_p("Ze względu na przyczepność:", indent=0.5, space_after=2)
    run(p, " c", bold=False)
    run(p, "min,b", bold=False, sub=True)

    p = add_p(indent=1.5)
    run(p, "Wymagane c")
    run(p, "min,b", sub=True)
    run(p, f" ≥ {inputs['fi_mm']} mm")
    if inputs["kruszywo_opis"] == "d_g > 32 mm":
        run(p, " (+5 mm ze względu na d")
        run(p, "g", sub=True)
        run(p, " > 32 mm)")

    p = add_p(indent=1.5)
    run(p, "Przyjęto: ")
    run(p, f"{wynik['c_min_b']:.0f} mm", bold=True)

    p = add_p("Ze względu na warunki środowiskowe:", indent=0.5, space_after=2)
    run(p, " c", bold=False)
    run(p, "min,dur", bold=False, sub=True)

    p = add_p(indent=1.5)
    run(p, f"Dla {inputs['klasa_ekspozycji']} i klasy {wynik['klasa_konstrukcji_final']} (Tablica 4.4N):")
    p = add_p(indent=1.5)
    run(p, "Wartość: ")
    run(p, f"{wynik['c_min_dur']:.0f} mm", bold=True)

    p = add_p("Całkowita otulina minimalna:", indent=0.5, space_after=2)
    run(p, " c", bold=False)
    run(p, "min", bold=False, sub=True)

    p = add_p(indent=1.5)
    run(p, "c")
    run(p, "min", sub=True)
    run(p, " = max( c")
    run(p, "min,b", sub=True)
    run(p, "; c")
    run(p, "min,dur", sub=True)
    run(p, " + Δc")
    run(p, "dur,γ", sub=True)
    run(p, " - Δc")
    run(p, "dur,st", sub=True)
    run(p, " - Δc")
    run(p, "dur,add", sub=True)
    run(p, "; 10 mm )")

    p = add_p(indent=1.5)
    run(p, "Wynik: ")
    run(p, "c")
    run(p, "min", sub=True)
    run(
        p,
        f" = max( {wynik['c_min_b']:.0f} mm; {wynik['c_min_dur']:.0f} mm + {inputs['dc_gamma']:.0f} mm - {inputs['dc_st']:.0f} mm - {inputs['dc_add']:.0f} mm; 10 mm )",
    )

    p = add_p(indent=1.5)
    run(p, " = ")
    run(p, f"{wynik['c_min']:.0f} mm", bold=True)

    # 4. Otulina nominalna
    doc.add_heading("4. Wyznaczenie otuliny nominalnej", level=1)

    p = add_p(indent=0.5)
    run(p, "c")
    run(p, "nom", sub=True)
    run(p, " = c")
    run(p, "min", sub=True)
    run(p, " + Δc")
    run(p, "dev", sub=True)

    p = add_p(indent=0.5)
    run(p, f"Przyjęta odchyłka Δc")
    run(p, "dev", sub=True)
    run(p, f" = {inputs['delta_dev']} mm")

    if wynik["limit_gruntu"] > 0:
        p = add_p(indent=0.5)
        run(p, f"Minimalna otulina na gruncie: {wynik['limit_gruntu']:.0f} mm")

    p = add_p(space_after=12)
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    
    run(p, "WYMAGANA OTULINA NOMINALNA: ", bold=False).font.size = Pt(12)
    
    r = run(p, "c")
    r.font.bold = True
    r.font.size = Pt(14)
    r = run(p, "nom")
    r.font.bold = True
    r.font.subscript = True
    r.font.size = Pt(14)
    r = run(p, f" = {wynik['c_nom']:.0f} mm", bold=True)
    r.font.size = Pt(14)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# STRONA GŁÓWNA STREAMLIT
# =============================================================================


def StronaOtulinaZbrojenia():
    st.markdown(
        """
        <style>
        .block-container { padding-top: 1.5rem; padding-bottom: 1.5rem; }
        h3 { margin-top: 1.0rem !important; margin-bottom: 0.4rem !important; font-size: 1.1rem; }
        
        /* USUNIĘTO AGRESYWNY STYL DLA PRZYCISKÓW */
        
        .big-result {
            font-size: 26px; font-weight: bold; color: #2E8B57; background-color: #f0f2f6;
            padding: 15px; border-radius: 8px; text-align: center; margin-top: 20px; border: 2px solid #2E8B57;
        }
        div.row-widget.stRadio > div { flex-direction: row; gap: 16px; }

        .header-help-icon {
            display:inline-flex;
            align-items:center;
            justify-content:center;
            width:18px;
            height:18px;
            border-radius:50%;
            border:1px solid #aaa;
            color:#aaa;
            font-size:11px;
            font-weight:600;
            cursor:help;
            transform:translateY(1px);
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Tytuł kalkulatora
    st.markdown(
        """
        <div style="text-align:center; margin-top:0.4rem; margin-bottom:0rem;">
            <span style="font-size:42px; font-weight:800; letter-spacing:1px; color:#dddddd;">
                OTULINA PRĘTÓW ZBROJENIOWYCH
            </span>
        </div>
        <div style="text-align:center; font-size:14px; color:#aaaaaa; margin-top:-12px; margin-bottom:0.6rem;">
            wg PN-EN 1992-1-1
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Dane wejściowe
    # REORDERED INPUTS: 1. Exposure, 2. Concrete, 3. Diameter
    st.markdown("### DANE WEJŚCIOWE")

    col_eksp, col_beton, col_fi = st.columns(3)
    
    with col_eksp:
        klasy_ekspozycji = [
            "X0",
            "XC1",
            "XC2",
            "XC3",
            "XC4",
            "XD1",
            "XD2",
            "XD3",
            "XS1",
            "XS2",
            "XS3",
        ]
        idx_eksp = klasy_ekspozycji.index("XC1")
        klasa_ekspozycji = st.selectbox(
            "Klasa ekspozycji", klasy_ekspozycji, index=idx_eksp
        )

    with col_beton:
        try:
            from TABLICE.ParametryBetonu import list_concrete_classes

            klasy_betonu = list_concrete_classes()
            idx_bet = klasy_betonu.index("C30/37") if "C30/37" in klasy_betonu else 0
        except ImportError:
            klasy_betonu = ["C25/30", "C30/37"]
            idx_bet = 1
        klasa_betonu = st.selectbox("Klasa betonu", klasy_betonu, index=idx_bet)

    with col_fi:
        try:
            from TABLICE.ParametryPretowZbrojeniowych import list_bar_diameters

            dostepne_fi = list_bar_diameters()
        except ImportError:
            dostepne_fi = [6, 8, 10, 12, 14, 16, 18, 20, 22, 25, 28, 32, 40]
        fi_mm = st.selectbox("Średnica pręta Φ [mm]", dostepne_fi, index=3)

    with st.expander("ℹ️ Pomoc: Opis klas ekspozycji (Tablica 4.1)", expanded=False):
        rys_eksp = SCIEZKA_FOLDERU_LOKALNEGO / "Otulina_klasy ekspozycji.png"
        if rys_eksp.exists():
            col_l, col_c, col_r = st.columns([1, 2, 1])
            with col_c:
                st.image(str(rys_eksp), use_container_width=True)
        else:
            st.info(f"Brak pliku pomocy: {rys_eksp.name}")
    st.markdown("---")

    # Uwarunkowania konstrukcyjne
    st.markdown("#### Uwarunkowania konstrukcyjne")
    col_okres, col_geom, col_kontrola = st.columns(3)

    with col_okres:
        st.markdown("**Projektowy okres użytkowania**")
        okres_uzytkowania = st.radio(
            "okres",
            ["50 lat", "100 lat"],
            index=0,
            label_visibility="collapsed",
            key="radio_okres",
        )
        is_100_lat = okres_uzytkowania == "100 lat"

    with col_geom:
        st.markdown("**Geometria elementu**")
        geometria = st.radio(
            "geometria",
            ["Płyta", "Belka / Słup / Inne"],
            index=0,
            label_visibility="collapsed",
            key="radio_geom",
        )
        is_plyta = geometria == "Płyta"

    with col_kontrola:
        st.markdown("**Specjalna kontrola jakości**")
        kontrola_jakosci_str = st.radio(
            "kontrola",
            ["Nie", "Tak"],
            index=0,
            label_visibility="collapsed",
            key="radio_kontrola",
        )
        kontrola = kontrola_jakosci_str == "Tak"

    with st.expander("ℹ️ Pomoc: Klasyfikacja konstrukcji (Tablica 4.3N)", expanded=False):
        rys_klasy = SCIEZKA_FOLDERU_LOKALNEGO / "Otulina_klasy konstrukcji.png"
        if rys_klasy.exists():
            col_l, col_c, col_r = st.columns([1, 2, 1])
            with col_c:
                st.image(str(rys_klasy), use_container_width=True)
        else:
            st.info(f"Brak pliku: {rys_klasy.name}")
    st.markdown("---")

    # Uwarunkowania wykonawcze (COMPRESSED - 3 COLUMNS)
    st.markdown("#### Uwarunkowania wykonawcze")
    
    # Initialization
    dc_gamma = 0.0
    dc_st = 0.0
    dc_add = 0.0

    c_grunt, c_krusz, c_odchylki_q = st.columns(3)

    with c_grunt:
        st.markdown("**Betonowanie na gruncie**")
        grunt_opts = ["Nie", "Na przygotowanym podłożu", "Bezpośrednio na gruncie"]
        betonowanie_grunt = st.radio(
            "grunt",
            grunt_opts,
            index=0,
            label_visibility="collapsed",
            key="radio_grunt",
        )

    with c_krusz:
        st.markdown("**Maksymalny wymiar kruszywa**")
        kruszywo_opts = ["$d_g \\le 32$ mm", "$d_g > 32$ mm"]
        kruszywo_str = st.radio(
            "kruszywo",
            kruszywo_opts,
            index=0,
            label_visibility="collapsed",
            key="radio_kruszywo",
        )
        dg_gt_32 = kruszywo_str == "$d_g > 32$ mm"
    kruszywo_opis = "d_g ≤ 32 mm" if not dg_gt_32 else "d_g > 32 mm"

    with c_odchylki_q:
        # Styled Header with Help Icon
        st.markdown(
            """
            <div style="display:flex;align-items:center;gap:6px;">
              <span style="font-weight:600; font-size:1rem;">Odchyłki dodatkowe</span>
              <span class="header-help-icon"
                    title="Wg Załącznika Krajowego do PN-EN 1992-1-1 zalecane wartości wynoszą: Δc_dur,γ = 0 mm, Δc_dur,st = 0 mm, Δc_dur,add = 0 mm.">?</span>
            </div>
            <div style="font-size: 13px; color: #888; margin-bottom: 5px;">
            (Δc<sub>dur,γ</sub>, Δc<sub>dur,st</sub>, Δc<sub>dur,add</sub>)
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Explicit Question Label (Cleaned)
        st.write("Czy przyjąć inne wartości niż zalecane w NA?")
        
        u_odchylki_str = st.radio(
            "odchylki_yn",
            ["Nie", "Tak"],
            index=0,
            label_visibility="collapsed",
            horizontal=True
        )
        
        if u_odchylki_str == "Tak":
            st.markdown("<div style='margin-top: 10px;'></div>", unsafe_allow_html=True)
            dc_gamma = st.number_input(r"$\Delta c_{dur,\gamma}$ [mm]", value=0.0, step=1.0)
            dc_st = st.number_input(r"$\Delta c_{dur,st}$ [mm]", value=0.0, step=1.0)
            dc_add = st.number_input(r"$\Delta c_{dur,add}$ [mm]", value=0.0, step=1.0)

    st.markdown("---")

    # Odchyłka wykonawcza
    st.markdown(
        """
        <div style="display:flex;align-items:center;gap:6px;">
          <h4 style="margin:0;">Odchyłka wykonawcza Δc<sub>dev</sub> (Wpływ na c<sub>nom</sub>)</h4>
          <span class="header-help-icon"
                title="Domyślnie 10 mm (EC2). Można zmniejszyć do 5 lub 0 mm przy wysokiej kontroli.">?</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    delta_dev = st.selectbox(
        "delta_dev_label",
        [0, 5, 10],
        index=2,
        key="delta_dev",
        label_visibility="collapsed"
    )

    st.markdown("---")

    # OBLICZENIA - DODANO TYPE PRIMARY
    _, c_center, _ = st.columns([1, 1.5, 1])
    with c_center:
        oblicz = st.button("OBLICZ OTULINĘ", type="primary", use_container_width=True)

    if oblicz:
        try:
            wynik = ObliczOtuline(
                klasa_ekspozycji=klasa_ekspozycji,
                fi_mm=float(fi_mm),
                klasa_betonu=klasa_betonu,
                zywotnosc_100_lat=is_100_lat,
                element_plytowy=is_plyta,
                kontrola_jakosci=kontrola,
                beton_na_gruncie=betonowanie_grunt,
                dg_gt_32=dg_gt_32,
                delta_dur_gamma=dc_gamma,
                delta_dur_st=dc_st,
                delta_dur_add=dc_add,
                delta_dev=float(delta_dev),
            )

            st.session_state["wynik_otuliny"] = wynik
            st.session_state["inputs_otuliny"] = {
                "klasa_betonu": klasa_betonu,
                "klasa_ekspozycji": klasa_ekspozycji,
                "fi_mm": fi_mm,
                "zywotnosc_100": is_100_lat,
                "plyta": is_plyta,
                "delta_dev": delta_dev,
                "betonowanie_grunt": betonowanie_grunt,
                "kontrola_jakosci_str": kontrola_jakosci_str,
                "kruszywo_opis": kruszywo_opis,
                "dc_gamma": dc_gamma,
                "dc_st": dc_st,
                "dc_add": dc_add,
            }
            st.session_state["pokaz_otuline"] = True

        except Exception as e:
            st.error(f"Wystąpił błąd podczas obliczeń: {e}")
            st.session_state["pokaz_otuline"] = False

    # WYNIKI
    if st.session_state.get("pokaz_otuline", False):
        wynik = st.session_state["wynik_otuliny"]
        inputs = st.session_state["inputs_otuliny"]
        c_nom = wynik["c_nom"]
        c_min = wynik["c_min"]

        st.markdown(
            f"""
            <div class="big-result">
                Nominalna otulina <i>c</i><sub>nom</sub> = {c_nom:.0f} mm
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown("<div style='height: 12px;'></div>", unsafe_allow_html=True)
        col_dl1, col_dl2 = st.columns(2)

        from_streamlit_pdf = create_pdf_report(wynik, inputs)
        from_streamlit_docx = create_docx_report(wynik, inputs)

        with col_dl1:
            st.download_button(
                "📄 POBIERZ RAPORT PDF",
                from_streamlit_pdf,
                "Raport_Otulina_EC2.pdf",
                "application/pdf",
                use_container_width=True,
            )

        with col_dl2:
            st.download_button(
                "📝 POBIERZ RAPORT WORD",
                from_streamlit_docx,
                "Raport_Otulina_EC2.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)

        with st.expander("Szczegóły obliczeń", expanded=False):
            st.write("**1. Klasa konstrukcji (Tablica 4.3N):**")

            zmiana = wynik["zmiana_klasy"]
            skladniki = []

            if inputs.get("zywotnosc_100"):
                skladniki.append(
                    (2, "Zwiększenie okresu użytkowania (100 lat)")
                )

            try:
                fck_local = int(
                    str(inputs["klasa_betonu"]).split("/")[0].replace("C", "")
                )
            except Exception:
                fck_local = 20

            eksp_local = wynik.get("klasa_ekspozycji", inputs["klasa_ekspozycji"])
            redukcja_wytrz = False
            if eksp_local == "XC1" and fck_local >= 30:
                redukcja_wytrz = True
            elif eksp_local in ["XC2", "XC3"] and fck_local >= 35:
                redukcja_wytrz = True
            elif eksp_local == "XC4" and fck_local >= 40:
                redukcja_wytrz = True
            elif eksp_local in ["XD1", "XD2", "XS1"] and fck_local >= 40:
                redukcja_wytrz = True
            elif eksp_local in ["XD3", "XS2", "XS3"] and fck_local >= 45:
                redukcja_wytrz = True

            if redukcja_wytrz:
                skladniki.append(
                    (-1, f"Redukcja ze względu na klasę wytrzymałości ({inputs['klasa_betonu']})")
                )

            if inputs.get("plyta") and eksp_local == "XC1":
                skladniki.append(
                    (-1, "Element płytowy w klasie ekspozycji XC1")
                )

            if inputs.get("kontrola_jakosci_str") == "Tak":
                skladniki.append(
                    (-1, "Specjalna kontrola jakości produkcji betonu")
                )

            st.markdown("Bazowa klasa konstrukcji: S4.")
            st.markdown("Modyfikacja klasy konstrukcji ze względu na warunki projektowe:")

            if not skladniki:
                st.markdown("&nbsp;&nbsp;&nbsp;&nbsp;- Brak modyfikacji", unsafe_allow_html=True)
            else:
                for delta, opis in skladniki:
                    znak = "+" if delta > 0 else ""
                    st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;ΔS = {znak}{delta}: {opis}", unsafe_allow_html=True)

            st.markdown(f"Obliczeniowa klasa konstrukcji: **{wynik['klasa_konstrukcji_final']}**")

            st.markdown("<br>", unsafe_allow_html=True)

            st.write("**2. Składowe $c_{min}$:**")
            st.latex(rf"c_{{min,b}} = {wynik['c_min_b']:.0f}\text{{ mm }}")
            st.latex(rf"c_{{min,dur}} = {wynik['c_min_dur']:.0f}\text{{ mm }}")
            st.latex(
                r"c_{min} = \max\Big(c_{min,b};\; c_{min,dur} + \Delta c_{dur,\gamma} - \Delta c_{dur,st} - \Delta c_{dur,add};\; 10\text{ mm}\Big)"
            )

            st.latex(
                rf"c_{{min}} = \max\Big({wynik['c_min_b']:.0f}\text{{ mm}};\; "
                rf"{wynik['c_min_dur']:.0f}\text{{ mm}} + {inputs['dc_gamma']:.0f}\text{{ mm}} - {inputs['dc_st']:.0f}\text{{ mm}} - {inputs['dc_add']:.0f}\text{{ mm}};\; "
                r"10\text{ mm}\Big)"
                rf" = \mathbf{{{c_min:.0f}\text{{ mm}}}}"
            )

            if wynik["limit_gruntu"] > 0:
                st.write("**3. Warunek gruntowy (betonowanie na gruncie):**")
                st.write(
                    f"- Sposób betonowania: **{inputs['betonowanie_grunt']}** \n"
                    f"- Wymagana minimalna otulina ze względu na warunki gruntowe: "
                    f"$c_{{nom}} \\ge {wynik['limit_gruntu']:.0f} \\text{{ mm}}$"
                )

            st.write("**4. Wynik końcowy:**")
            st.latex("c_{nom} = c_{min} + \\Delta c_{dev}")
            st.latex(
                rf"c_{{nom}} = {c_min:.0f}\text{{ mm}} + "
                rf"{float(inputs['delta_dev']):.0f}\text{{ mm}}"
                rf" = \mathbf{{{c_nom:.0f}\text{{ mm}}}}"
            )


if __name__ == "__main__":
    StronaOtulinaZbrojenia()