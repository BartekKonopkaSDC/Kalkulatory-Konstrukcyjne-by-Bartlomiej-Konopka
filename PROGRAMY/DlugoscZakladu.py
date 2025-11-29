"""
PROGRAMY/DlugoscZakladu.py
Wersja: ENGINEERING_REPORT_V20 (FULL).
"""

from __future__ import annotations

import streamlit as st
from pathlib import Path
from io import BytesIO

# Biblioteki do PDF
from fpdf import FPDF

# Biblioteki do DOCX
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from TABLICE.ParametryPretowZbrojeniowych import get_bar_params
from OBLICZENIA.Obliczenia_DlugoscZakladu import ObliczDlugoscZakladu
from TABLICE.ParametryBetonu import get_concrete_params, list_concrete_classes
from TABLICE.ParametryStali import get_steel_params, list_steel_grades

SCIEZKA_BAZOWA = Path(__file__).resolve().parents[1]
SCIEZKA_DODATKI = SCIEZKA_BAZOWA / "DODATKI"
SCIEZKA_PROGRAMY = Path(__file__).parent 

# --- SYMBOLE UNICODE DLA PDF ---
SYM = {
    "fi": "\u03A6", "alpha": "\u03B1", "sigma": "\u03C3", "eta": "\u03B7",
    "rho": "\u03C1", "dot": "\u00B7", "bullet": "\u2022", "ge": "\u2265", "le": "\u2264",
    "ra": "\u2192"
}

# =============================================================================
# GENERATOR PDF
# =============================================================================

def create_pdf_report(wynik: dict, inputs: dict) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    
    font_regular = SCIEZKA_PROGRAMY / "DejaVuSans.ttf"
    font_bold = SCIEZKA_PROGRAMY / "DejaVuSans-Bold.ttf"
    
    if font_regular.exists():
        pdf.add_font('DejaVu', '', str(font_regular), uni=True)
        if font_bold.exists():
            pdf.add_font('DejaVu', 'B', str(font_bold), uni=True)
        else:
            pdf.add_font('DejaVu', 'B', str(font_regular), uni=True)
        pdf.add_font('DejaVu', 'I', str(font_regular), uni=True)
        main_font = 'DejaVu'
    else:
        main_font = 'Arial'

    LINE_H = 6
    MARGIN_LEFT = 15
    MARGIN_TOP = 15
    X_INDENT = 20
    X_EQ = 25 

    pdf.set_margins(MARGIN_LEFT, MARGIN_TOP, 10)
    pdf.set_auto_page_break(True, margin=15)

    def w_txt(text, bold=False, italic=False, size=11):
        style = ''
        if bold: style += 'B'
        if italic: style += 'I'
        pdf.set_font(main_font, style, size)
        pdf.write(LINE_H, text)

    def w_sub(text, size=7, move_down=2.0):
        orig_y, orig_x = pdf.get_y(), pdf.get_x()
        pdf.set_font(main_font, '', size)
        pdf.set_xy(orig_x, orig_y + move_down)
        pdf.write(LINE_H, text)
        pdf.set_xy(pdf.get_x(), orig_y) 
        pdf.set_font(main_font, '', 11)

    def new_line():
        pdf.ln(LINE_H)

    def header_sec(text):
        pdf.ln(4)
        pdf.set_font(main_font, 'B', 12)
        pdf.cell(0, 8, text, ln=True, border='B')
        pdf.ln(2)

    def build_line(segments, indent=0):
        if indent > 0:
            pdf.set_x(indent)
        for item in segments:
            text = item[0]
            type_ = item[1]
            if type_ == 'txt': w_txt(text)
            elif type_ == 'bold': w_txt(text, bold=True)
            elif type_ == 'italic': w_txt(text, italic=True)
            elif type_ == 'sub': w_sub(text)
            elif type_ == 'sym': w_txt(text)
        new_line()

    # --- NAGŁÓWEK ---
    pdf.set_font(main_font, 'B', 14)
    pdf.cell(0, 8, "OBLICZENIE DŁUGOŚCI ZAKŁADU PRĘTÓW ZBROJENIOWYCH", ln=True, align='C')
    pdf.set_font(main_font, '', 10)
    pdf.cell(0, 5, "wg PN-EN 1992-1-1", ln=True, align='C')
    pdf.line(MARGIN_LEFT, pdf.get_y()+2, 200, pdf.get_y()+2)
    pdf.ln(5)

    # 1. MATERIAŁY
    header_sec("1. Parametry materiałowe")
    build_line([("Średnica pręta: ", 'txt'), (f"{SYM['fi']} = {wynik['fi_mm']} mm", 'txt')], X_INDENT)
    build_line([(f"Beton: {wynik['klasa_betonu']} (f", 'txt'), ("ctd", 'sub'), (f" = {wynik['fctd_MPa']:.2f} MPa)", 'txt')], X_INDENT)
    build_line([(f"Stal: {inputs['stal_nazwa']} (f", 'txt'), ("yk", 'sub'), (f" = {wynik['fyk_MPa']:.0f} MPa, f", 'txt'), ("yd", 'sub'), (f" = {wynik['f_yd_MPa']:.1f} MPa)", 'txt')], X_INDENT)
    build_line([("Naprężenia: ", 'txt'), (SYM['sigma'], 'sym'), ("sd", 'sub'), (f" = {wynik['sigma_sd_MPa']:.1f} MPa", 'txt')], X_INDENT)

    # 2. ZAKOTWIENIE
    header_sec("2. Podstawowa długość zakotwienia")
    build_line([(f"Warunki przyczepności: {inputs['warunki']} (", 'txt'), (SYM['eta'], 'sym'), ("1", 'sub'), (f" = {wynik['eta1']})", 'txt')], X_INDENT)
    build_line([(f"Współczynnik średnicy: ", 'txt'), (SYM['eta'], 'sym'), ("2", 'sub'), (f" = {wynik['eta2']:.2f}", 'txt')], X_INDENT)
    
    pdf.ln(1)
    
    # fbd
    build_line([("f", 'italic'), ("bd", 'sub'), (" = ", 'txt'), (f"2.25 {SYM['dot']} ", 'txt'), (SYM['eta'], 'sym'), ("1", 'sub'), (f" {SYM['dot']} ", 'txt'), (SYM['eta'], 'sym'), ("2", 'sub'), (f" {SYM['dot']} f", 'txt'), ("ctd", 'sub'), (f" = 2.25 {SYM['dot']} {wynik['eta1']} {SYM['dot']} {wynik['eta2']:.2f} {SYM['dot']} {wynik['fctd_MPa']:.2f} = ", 'txt'), (f"{wynik['f_bd_MPa']:.2f} MPa", 'txt')], X_EQ)
    
    # lb,rqd
    build_line([("l", 'italic'), ("b,rqd", 'sub'), (" = ", 'txt'), ("(", 'txt'), (SYM['fi'], 'txt'), (" / 4) ", 'txt'), (f"{SYM['dot']} ", 'txt'), ("(", 'txt'), (SYM['sigma'], 'sym'), ("sd", 'sub'), (" / f", 'txt'), ("bd", 'sub'), (")", 'txt'), (" = ", 'txt'), ("(", 'txt'), (str(wynik['fi_mm']), 'txt'), (" / 4) ", 'txt'), (f"{SYM['dot']} ", 'txt'), ("(", 'txt'), (f"{wynik['sigma_sd_MPa']:.1f}", 'txt'), (" / ", 'txt'), (f"{wynik['f_bd_MPa']:.2f}", 'txt'), (")", 'txt'), (" = ", 'txt'), (f"{wynik['l_b_rqd_mm']:.1f} mm", 'txt')], X_EQ)

    # 3. WSPÓŁCZYNNIKI
    header_sec(f"3. Współczynniki wpływu {SYM['alpha']}")
    indices = [1, 2, 3, 5, 6]
    for i in indices:
        val = wynik[f'alfa{i}']
        desc = ""
        if i == 6: desc = f"   ({SYM['rho']}"
        segments = [(SYM['alpha'], 'italic'), (str(i), 'sub'), (f" = {val:.2f}", 'txt')]
        if desc: segments.extend([(desc, 'txt'), ("1", 'sub'), (f" = {inputs['procent_lacz']}%)", 'txt')])
        build_line(segments, X_EQ)

    # WARUNEK 0.7
    pdf.ln(1)
    if wynik.get('limit_active', False):
        txt_warunek = f"{wynik['iloczyn_235_raw']:.2f} < 0.7 {SYM['ra']} przyjeto 0.7"
        build_line([("Warunek 8.4.4(1): ", 'bold'), (SYM['alpha'], 'italic'), ("2", 'sub'), (f"{SYM['dot']}", 'txt'), (SYM['alpha'], 'italic'), ("3", 'sub'), (f"{SYM['dot']}", 'txt'), (SYM['alpha'], 'italic'), ("5", 'sub'), (f" = {txt_warunek}", 'txt')], X_EQ)
    else:
        txt_warunek = f"{wynik['iloczyn_235_raw']:.2f} {SYM['ge']} 0.7 (OK)"
        build_line([("Warunek 8.4.4(1): ", 'txt'), (SYM['alpha'], 'italic'), ("2", 'sub'), (f"{SYM['dot']}", 'txt'), (SYM['alpha'], 'italic'), ("3", 'sub'), (f"{SYM['dot']}", 'txt'), (SYM['alpha'], 'italic'), ("5", 'sub'), (f" = {txt_warunek}", 'txt')], X_EQ)

    pdf.ln(2)
    build_line([(SYM['alpha'], 'italic'), ("glob", 'sub'), (" = ", 'txt'), (SYM['alpha'], 'italic'), ("1", 'sub'), (f" {SYM['dot']} ", 'txt'), (SYM['alpha'], 'italic'), ("2", 'sub'), (f" {SYM['dot']} ", 'txt'), (SYM['alpha'], 'italic'), ("3", 'sub'), (f" {SYM['dot']} ", 'txt'), (SYM['alpha'], 'italic'), ("5", 'sub'), (f" {SYM['dot']} ", 'txt'), (SYM['alpha'], 'italic'), ("6", 'sub'), (" = ", 'txt'), (f"{wynik['alfa']:.3f}", 'txt')], X_EQ)

    # 4. MINIMALNA
    header_sec("4. Minimalna długość zakładu")
    val1, val2, val3 = 0.3 * wynik['alfa6'] * wynik['l_b_rqd_mm'], 15 * wynik['fi_mm'], 200
    build_line([("l", 'italic'), ("0,min", 'sub'), (f" = max(0.3{SYM['alpha']}", 'txt'), ("6", 'sub'), ("l", 'txt'), ("b,rqd", 'sub'), (f"; 15{SYM['fi']}; 200 mm)", 'txt')], X_EQ)
    build_line([("      = max(", 'txt'), (f"{val1:.1f} mm; {val2:.1f} mm; {val3} mm", 'txt'), (") = ", 'txt'), (f"{wynik['L0_min_mm']:.1f} mm", 'txt')], X_EQ)

    # 5. WYNIK
    header_sec("5. Obliczenie długości zakładu")
    build_line([("l", 'italic'), ("0", 'sub'), (f" = {SYM['alpha']}", 'txt'), ("glob", 'sub'), (f" {SYM['dot']} l", 'txt'), ("b,rqd", 'sub'), (f" = {wynik['alfa']:.3f} {SYM['dot']} {wynik['l_b_rqd_mm']:.1f} = ", 'txt'), (f"{wynik['L0_mm']:.1f} mm", 'txt')], X_EQ)
    pdf.ln(1)
    build_line([("l", 'italic'), ("0,req", 'sub'), (" = max(l", 'txt'), ("0", 'sub'), ("; l", 'txt'), ("0,min", 'sub'), (") ", 'txt')], X_EQ)
    pdf.ln(4)
    pdf.set_fill_color(240, 240, 240)
    pdf.rect(MARGIN_LEFT, pdf.get_y(), 180, 12, 'F')
    pdf.set_y(pdf.get_y() + 3)
    
    # POGRUBIENIE WYNIKU CZCIONKĄ BOLD
    build_line([("Wymagana długość zakładu:  l", 'txt'), ("0,req", 'sub'), (f" = {wynik['L_z_mm']:.1f} mm", 'bold')], MARGIN_LEFT + 5)

    return pdf.output(dest='S').encode('latin-1', 'replace')


# =============================================================================
# GENERATOR DOCX
# =============================================================================

def create_docx_report(wynik: dict, inputs: dict) -> BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    if 'Heading 1' in doc.styles:
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)

    def add_p(indent=0, space_after=2):
        p = doc.add_paragraph()
        if indent > 0: p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def run_txt(p, text, bold=False, italic=False):
        r = p.add_run(text)
        r.font.bold = bold; r.font.italic = italic
        return r

    def run_sub(p, text, bold=False):
        r = p.add_run(text)
        r.font.subscript = True; r.font.bold = bold
        return r

    p = add_p(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_txt(p, "OBLICZENIE DŁUGOŚCI ZAKŁADU PRĘTÓW ZBROJENIOWYCH", bold=True).font.size = Pt(16)
    p = add_p(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_txt(p, "wg PN-EN 1992-1-1")
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Parametry materiałowe", level=1)
    p = add_p(1.0); run_txt(p, f"Średnica pręta: {SYM['fi']} = {wynik['fi_mm']} mm")
    p = add_p(1.0); run_txt(p, f"Beton: {wynik['klasa_betonu']} (f"); run_sub(p, "ctd"); run_txt(p, f" = {wynik['fctd_MPa']:.2f} MPa)")
    p = add_p(1.0); run_txt(p, f"Stal: {inputs['stal_nazwa']} (f"); run_sub(p, "yk"); run_txt(p, f" = {wynik['fyk_MPa']:.0f} MPa)")

    doc.add_heading("2. Podstawowa długość zakotwienia", level=1)
    p = add_p(1.0); run_txt(p, f"Warunki przyczepności: {inputs['warunki']}")
    # ZMIANA: Dodano eta2
    p = add_p(1.0); run_txt(p, "Współczynnik średnicy: "); run_txt(p, SYM['eta'], italic=True); run_sub(p, "2"); run_txt(p, f" = {wynik['eta2']:.2f}")
    
    p = add_p(1.5, space_after=6)
    run_txt(p, "f", italic=True); run_sub(p, "bd"); run_txt(p, f" = 2.25 {SYM['dot']} {wynik['eta1']} {SYM['dot']} {wynik['eta2']:.2f} {SYM['dot']} f"); run_sub(p, "ctd"); run_txt(p, f" = {wynik['f_bd_MPa']:.2f} MPa")
    
    p = add_p(1.5, space_after=6)
    run_txt(p, "l", italic=True); run_sub(p, "b,rqd"); run_txt(p, f" = ({wynik['fi_mm']}/4) {SYM['dot']} ("); run_txt(p, SYM['sigma'], italic=True); run_sub(p, "sd"); run_txt(p, "/f"); run_sub(p, "bd"); run_txt(p, f") = {wynik['l_b_rqd_mm']:.1f} mm")

    doc.add_heading(f"3. Współczynniki wpływu {SYM['alpha']}", level=1)
    for i in [1, 2, 3, 5, 6]:
        p = add_p(1.5, space_after=0)
        run_txt(p, SYM['alpha'], italic=True); run_sub(p, str(i)); run_txt(p, f" = {wynik[f'alfa{i}']:.2f}")
        if i == 6: run_txt(p, f" ({SYM['rho']}"); run_sub(p, "1"); run_txt(p, f" = {inputs['procent_lacz']}%)")
    p = add_p(1.5, space_after=6)
    p.paragraph_format.space_before = Pt(6)
    run_txt(p, SYM['alpha'], italic=True); run_sub(p, "glob"); run_txt(p, " = "); run_txt(p, SYM['alpha'], italic=True); run_sub(p, "1"); run_txt(p, f" {SYM['dot']} "); run_txt(p, SYM['alpha'], italic=True); run_sub(p, "2"); run_txt(p, f" {SYM['dot']} "); run_txt(p, SYM['alpha'], italic=True); run_sub(p, "3"); run_txt(p, f" {SYM['dot']} "); run_txt(p, SYM['alpha'], italic=True); run_sub(p, "5"); run_txt(p, f" {SYM['dot']} "); run_txt(p, SYM['alpha'], italic=True); run_sub(p, "6"); run_txt(p, f" = {wynik['alfa']:.3f}")

    doc.add_heading("4. Minimalna długość zakładu", level=1)
    p = add_p(1.5, space_after=6)
    run_txt(p, "l", italic=True); run_sub(p, "0,min"); run_txt(p, f" = max(0.3{SYM['alpha']}"); run_sub(p, "6"); run_txt(p, "l"); run_sub(p, "b,rqd"); run_txt(p, f"; 15{SYM['fi']}; 200 mm) = "); run_txt(p, f"{wynik['L0_min_mm']:.1f} mm")

    doc.add_heading("5. Obliczenie długości zakładu", level=1)
    p = add_p(1.5, space_after=6)
    run_txt(p, "l", italic=True); run_sub(p, "0"); run_txt(p, f" = {SYM['alpha']}"); run_sub(p, "glob"); run_txt(p, f" {SYM['dot']} l"); run_sub(p, "b,rqd"); run_txt(p, f" = {wynik['L0_mm']:.1f} mm")
    p = add_p(1.0, space_after=12)
    p.paragraph_format.space_before = Pt(12)
    run_txt(p, "Wymagana długość zakładu l"); run_sub(p, "0,req"); run_txt(p, f" = {wynik['L_z_mm']:.1f} mm", bold=True)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# STRONA GŁÓWNA STREAMLIT
# =============================================================================

def StronaDlugoscZakladu() -> None:
    st.markdown("""
    <style>
    .block-container { padding-top: 3rem; padding-bottom: 1rem; }
    h1, h2, h3, h4, h5 { margin-bottom: 0.2rem !important; margin-top: 0.5rem !important; }
    div.row-widget.stRadio > div { flex-direction: row; align-items: center; }
    div.row-widget.stRadio > div > label { margin-right: 15px; }
    div.stButton > button:first-child {
        background-color: #0099ff; color: white; font-size: 20px; font-weight: 600;
        width: 100%; border-radius: 8px; padding: 0.75rem 1rem; margin-top: 1rem;
    }
    .big-result {
        font-size: 28px; font-weight: bold; color: #2E8B57; background-color: #f0f2f6;
        padding: 15px; border-radius: 10px; text-align: center; margin-top: 10px; border: 2px solid #2E8B57;
    }
    </style>
    """, unsafe_allow_html=True)

    st.title("DŁUGOŚĆ ZAKŁADU PRĘTÓW wg PN-EN 1992-1-1")

    col_main, _ = st.columns([1.0, 0.01])

    with col_main:
        st.subheader("1. DANE WEJŚCIOWE")
        
        c1, c2, c3 = st.columns(3)
        with c1:
            dostepne_fi = [6, 8, 10, 12, 14, 16, 18, 20, 22, 25, 28, 32, 40]
            fi_mm = st.selectbox("Średnica pręta Φ [mm]", dostepne_fi, index=3)
        with c2:
            klasy_betonu = list_concrete_classes()
            idx_bet = klasy_betonu.index("C30/37") if "C30/37" in klasy_betonu else 0
            klasa_betonu = st.selectbox("Klasa betonu", klasy_betonu, index=idx_bet)
        with c3:
            dostepne_stale = list_steel_grades()
            wybrana_stal = st.selectbox("Klasa stali", dostepne_stale, index=0)

        As_bar_mm2 = get_bar_params(fi_mm).As
        stal = get_steel_params(wybrana_stal)
        fyk_MPa = stal.fyk 

        c_napr, c_war, c_rys = st.columns([1.6, 0.7, 2.7])
        with c_napr:
            st.markdown("**Naprężenia w stali $\\sigma_{sd} / f_{yd}$ [%]**")
            naprezenia_stali = st.number_input(" ", min_value=0.0, max_value=100.0, value=100.0, step=1.0, format="%.0f", key="naprezenia_input", label_visibility="collapsed")
        with c_war:
            st.markdown("**Warunki przyczepności**")
            warunki = st.radio(" ", ["Dobre", "Złe"], horizontal=True, label_visibility="collapsed")
            eta1 = 1.0 if warunki == "Dobre" else 0.7
        with c_rys:
            st.markdown("&nbsp;", unsafe_allow_html=True) 
            with st.expander("Rysunek pomocniczy: Warunki przyczepności"):
                rys = SCIEZKA_DODATKI / "DlugoscZakotwieniaIZakladu_WarunkiPrzyczepnosci.png"
                if rys.exists(): st.image(rys, use_container_width=True)

        st.markdown("---")
        st.subheader("2. WSPÓŁCZYNNIKI WPŁYWU $\\alpha$")
        
        c_rodzaj, c_rys_alfa = st.columns([1, 1.5])
        with c_rodzaj:
            rodzaj_preta = st.radio("Rodzaj pręta:", ["Rozciągany", "Ściskany"], index=1, horizontal=True)
        with c_rys_alfa:
            with st.expander("Rysunek pomocniczy: Współczynniki wpływu alfa"):
                rys_a15 = SCIEZKA_DODATKI / "DlugoscZakotwieniaIZakladu_alfa1-alfa5.png"
                rys_a6 = SCIEZKA_DODATKI / "DlugoscZakladu_alfa6.png"
                if rys_a15.exists(): st.image(rys_a15, use_container_width=True, caption="Tablica 8.2")
                if rys_a6.exists(): st.image(rys_a6, use_container_width=True, caption="Tablica 8.3")

        alfa1, alfa2, alfa3, alfa4, alfa5, alfa6 = 1.0, 1.0, 1.0, 1.0, 1.0, 1.0
        wsp_cd = 30.0

        if rodzaj_preta == "Rozciągany":
            col_a2, col_a3, col_a5 = st.columns(3)
            with col_a2:
                st.markdown("**$\\alpha_2$: Wpływ otuliny ($c_d$)**")
                u_cd = st.checkbox("Czy uwzględnić wpływ otuliny?", value=False)
                if u_cd:
                    with st.expander("Rysunek $c_d$"):
                        rys_cd = SCIEZKA_DODATKI / "DlugoscZakotwieniaIZakladu_Wspolczynnik cd.png"
                        if rys_cd.exists(): st.image(rys_cd, use_container_width=True)
                    wsp_cd = st.number_input("Współczynnik $c_{d}$ [mm]", min_value=0.0, value=30.0, step=1.0)
                    if fi_mm > 0:
                        alfa2_raw = 1.0 - 0.15 * ((wsp_cd - fi_mm) / fi_mm)
                        alfa2 = max(0.7, min(1.0, alfa2_raw))
            with col_a3:
                st.markdown("**$\\alpha_3$: Zbrojenie poprzeczne**")
                u_a3 = st.checkbox("Czy uwzględnić zbrojenie poprzeczne?", value=False)
                if u_a3:
                    with st.expander("Rysunek $K$"):
                        rys_k = SCIEZKA_DODATKI / "DlugoscZakotwieniaIZakladu_Wspolczynnik K.png"
                        if rys_k.exists(): st.image(rys_k, use_container_width=True)
                    wsp_K = st.selectbox("Współczynnik $K$", [0.0, 0.05, 0.1], index=1)
                    sum_Ast_cm2 = st.number_input("$\\Sigma A_{st}$ [cm²]", value=0.0, step=0.01)
                    sum_Ast_min_cm2 = st.number_input("$\\Sigma A_{st,min}$ [cm²]", value=2.50, step=0.01)
                    s_ast_mm = sum_Ast_cm2 * 100
                    s_ast_min_mm = sum_Ast_min_cm2 * 100
                    if As_bar_mm2 > 0:
                        lambda_val = (s_ast_mm - s_ast_min_mm) / As_bar_mm2
                        lambda_val = max(0.0, lambda_val)
                        a3_raw = 1.0 - wsp_K * lambda_val
                        alfa3 = max(0.7, min(1.0, a3_raw))
            with col_a5:
                st.markdown("**$\\alpha_5$: Nacisk poprzeczny**")
                u_a5 = st.checkbox("Czy uwzględnić nacisk poprzeczny?", value=False)
                if u_a5:
                    p = st.number_input("p - nacisk poprzeczny w [MPa] wzdłuż lbd w stanie granicznym nośności", value=8.0, step=0.1)
                    alfa5 = max(0.7, min(1.0, 1.0 - 0.04 * p))
        
        st.write("") 
        c_label, c_input = st.columns([2, 1])
        with c_label:
            st.markdown("**$\\alpha_6$: Procent prętów łączonych ($\\rho_1$)**")
        with c_input:
            procent_lacz = st.selectbox(" ", [100, 50, 33, 25], index=1, label_visibility="collapsed")
        
        mapa_alfa6 = {100: 1.50, 50: 1.4, 33: 1.15, 25: 1.00}
        alfa6 = mapa_alfa6.get(procent_lacz, 1.5)

        st.markdown("---")

        # --- OBLICZENIA ---
        c_left, c_center, c_right = st.columns([1, 1.5, 1]) 
        with c_center:
            oblicz = st.button("OBLICZ DŁUGOŚĆ ZAKŁADU", use_container_width=True)

        if oblicz:
            wynik = ObliczDlugoscZakladu(
                fi_mm=fi_mm, klasa_betonu=klasa_betonu, fyk_MPa=float(fyk_MPa),
                procent_naprezenia=float(naprezenia_stali), eta1=eta1, wsp_cd=wsp_cd,
                alfa1=alfa1, alfa2=alfa2, alfa3=alfa3, alfa4=alfa4, alfa5=alfa5, alfa6=alfa6,
            )
            
            inputs_desc = {
                "stal_nazwa": wybrana_stal,
                "warunki": warunki,
                "rodzaj_preta": rodzaj_preta,
                "procent_lacz": procent_lacz
            }
            st.session_state['wynik_zakladu'] = wynik
            st.session_state['inputs_zakladu'] = inputs_desc
            st.session_state['pokaz_wyniki'] = True

        # --- WYNIKI ---
        if st.session_state.get('pokaz_wyniki'):
            wynik = st.session_state['wynik_zakladu']
            inputs_pdf = st.session_state['inputs_zakladu']
            
            L_z_mm = wynik["L_z_mm"]
            f_ctd_val = wynik["fctd_MPa"]
            f_bd_val = wynik["f_bd_MPa"]
            sigma_sd_val = wynik["sigma_sd_MPa"]
            l_b_rqd_mm = wynik["l_b_rqd_mm"]
            L0_mm = wynik["L0_mm"]
            L0_min_mm = wynik["L0_min_mm"]
            alfa_glob = wynik["alfa"]

            # DUŻY WYNIK
            st.markdown(f"""
            <div class="big-result">
                Wymagana długość zakładu <i>l</i><sub>0,req</sub> = {L_z_mm:.0f} mm
            </div>
            """, unsafe_allow_html=True)

            # --- PRZYCISKI POBIERANIA ---
            st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
            col_dl1, col_dl2 = st.columns(2)
            
            with col_dl1:
                pdf_bytes = create_pdf_report(wynik, inputs_pdf)
                st.download_button(
                    label="📄 POBIERZ RAPORT PDF",
                    data=pdf_bytes,
                    file_name="Raport_Dlugosc_Zakladu_EC2.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

            with col_dl2:
                docx_buffer = create_docx_report(wynik, inputs_pdf)
                st.download_button(
                    label="📝 POBIERZ RAPORT WORD",
                    data=docx_buffer,
                    file_name="Raport_Dlugosc_Zakladu_EC2.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)

            # --- SZCZEGÓŁY ---
            with st.expander("📄 Szczegółowy raport z obliczeń", expanded=False):
                st.markdown("#### 1. Parametry materiałowe")
                st.write(f"- Średnica pręta: **Φ = {fi_mm} mm**")
                st.write(f"- Beton: **{klasa_betonu}** ($f_{{ctd}} = {f_ctd_val:.2f}$ MPa)")
                st.write(f"- Stal: **{wybrana_stal}** ($f_{{yk}} = {wynik['fyk_MPa']:.0f}$ MPa, $f_{{yd}} = {wynik['f_yd_MPa']:.1f}$ MPa)")
                st.write(f"- Naprężenia obliczeniowe: $\sigma_{{sd}} = {sigma_sd_val:.1f}$ MPa ({naprezenia_stali:.0f}% $f_{{yd}}$)")
                
                st.markdown("#### 2. Podstawowa długość zakotwienia ($l_{b,rqd}$)")
                st.write(f"- Warunki przyczepności: **{warunki}** ($\eta_1 = {eta1}$)")
                st.write(f"- Współczynnik średnicy: $\eta_2 = {wynik['eta2']:.2f}$") # ZMIANA: Dodano eta2
                st.write(f"- Przyczepność graniczna $f_{{bd}}$:")
                # ZMIANA: Dodano eta2 do wzoru
                st.latex(rf"f_{{bd}} = 2.25 \cdot \eta_1 \cdot \eta_2 \cdot f_{{ctd}} = 2.25 \cdot {eta1} \cdot {wynik['eta2']} \cdot {f_ctd_val:.2f} = \mathbf{{{f_bd_val:.2f} \text{{ MPa}}}}")
                st.latex(rf"l_{{b,rqd}} = \frac{{\Phi}}{{4}} \cdot \frac{{\sigma_{{sd}}}}{{f_{{bd}}}} = \frac{{{fi_mm}}}{{4}} \cdot \frac{{{sigma_sd_val:.1f}}}{{{f_bd_val:.2f}}} = \mathbf{{{l_b_rqd_mm:.1f} \text{{ mm}}}}")

                st.markdown("#### 3. Współczynniki wpływu $\\alpha$")
                st.write(f"- $\\alpha_1 = {wynik['alfa1']:.2f}$")
                st.write(f"- $\\alpha_2 = {wynik['alfa2']:.2f}$")
                st.write(f"- $\\alpha_3 = {wynik['alfa3']:.2f}$")
                st.write(f"- $\\alpha_5 = {wynik['alfa5']:.2f}$")
                st.write(f"- $\\alpha_6 = {wynik['alfa6']:.2f}$ ($\\rho_1 = {procent_lacz}\%$)")
                
                # --- NOWOŚĆ: Wyświetlanie limitu 0.7 ---
                st.write("---")
                iloczyn = wynik['iloczyn_235_raw']
                if wynik['limit_active']:
                    st.error(f"⚠️ Warunek EC2 8.4.4(1): $\\alpha_2 \\cdot \\alpha_3 \\cdot \\alpha_5 = {iloczyn:.3f} < 0.7$")
                    st.latex(r"\rightarrow \text{Przyjęto: } \alpha_2 \cdot \alpha_3 \cdot \alpha_5 = 0.7")
                else:
                    st.success(f"✅ Warunek EC2 8.4.4(1): $\\alpha_2 \\cdot \\alpha_3 \\cdot \\alpha_5 = {iloczyn:.3f} \ge 0.7$")
                st.write("---")

                st.latex(rf"\alpha_{{global}} = \alpha_1 \cdot \alpha_2 \cdot \alpha_3 \cdot \alpha_5 \cdot \alpha_6 = \mathbf{{{alfa_glob:.3f}}}")

                st.markdown("#### 4. Minimalna długość zakładu ($l_{0,min}$)")
                val_min1 = 0.3 * alfa6 * l_b_rqd_mm
                val_min2 = 15 * fi_mm
                st.latex(rf"l_{{0,min}} = \max(0.3 \cdot \alpha_6 \cdot l_{{b,rqd}}; 15\Phi; 200) = \max({val_min1:.1f}; {val_min2:.1f}; 200) = {L0_min_mm:.1f} \text{{ mm}}")

                st.markdown("#### 5. Obliczenie długości zakładu ($l_0$)")
                st.latex(rf"l_0 = \alpha_{{global}} \cdot l_{{b,rqd}} = {alfa_glob:.3f} \cdot {l_b_rqd_mm:.1f} = {L0_mm:.1f} \text{{ mm}}")
                st.latex(rf"l_{{0,req}} = \max(l_0; l_{{0,min}}) = \mathbf{{{L_z_mm:.1f} \text{{ mm}}}}")

    with st.expander("Uwagi / założenia obliczeń"):
        st.markdown("- Obliczenia dotyczą zakładu **prostych prętów żebrowanych** wg PN-EN 1992-1-1 (EC2).")

if __name__ == "__main__":
    StronaDlugoscZakladu()